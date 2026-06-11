import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter import Toplevel, Entry, Button, Frame, Label, Canvas, Scrollbar, BOTH, LEFT, RIGHT, Y, X, BOTTOM, W, E, NORMAL, DISABLED, Checkbutton, IntVar
import pandas as pd
import sqlite3
import xml.etree.ElementTree as ET
from openpyxl.utils import get_column_letter
import os
import traceback
import datetime
import json
import configparser  # ⭐ THÊM IMPORT NÀY

# Thư viện để xuất Word
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ======================= CẤU HÌNH ĐƯỜNG DẪN =======================

DB_FILE = "nguoi_lx.db"
LOG_DIR = "logs"
CONFIG_DIR = os.path.dirname(os.path.abspath(__file__))
EXCEL_CONFIG_FILE = os.path.join(CONFIG_DIR, "config_excel.json")
NOIDUNG_CONFIG_FILE = os.path.join(CONFIG_DIR, "config_noidung.json")
DISPLAY_CONFIG_FILE = os.path.join(CONFIG_DIR, "config_display.json")
DM_DVHC_FILE = os.path.join(CONFIG_DIR, "DM_DVHC.cfg")  # ⭐ FILE CFG

os.makedirs(LOG_DIR, exist_ok=True)


# ======================= BẢNG NỘI DUNG THI =======================
# True = cần thi, False = không thi
NOIDUNG_MAP = {
    "1":  {"LT": True,  "MP": False, "H": True,  "D": False},  # LT+H lần đầu A1
    "2":  {"LT": True,  "MP": False, "H": False, "D": False},  # LT
    "3":  {"LT": True,  "MP": False, "H": True,  "D": False},  # LT+H
    "4":  {"LT": True,  "MP": False, "H": True,  "D": True},   # LT+H+D
    "5":  {"LT": False, "MP": False, "H": True,  "D": True},   # H+D
    "6":  {"LT": False, "MP": False, "H": False, "D": True},   # D
    "7":  {"LT": False, "MP": False, "H": True,  "D": False},  # H
    "8":  {"LT": False, "MP": False, "H": True,  "D": False},  # H miễn LT ô tô
    "9":  {"LT": False, "MP": False, "H": True,  "D": False},  # H miễn LT xe máy
    "11": {"LT": True,  "MP": True,  "H": True,  "D": True},   # LT+MP+H+D lần đầu
    "12": {"LT": True,  "MP": True,  "H": True,  "D": True},   # LT+MP+H+D lại
    "13": {"LT": False, "MP": True,  "H": True,  "D": True},   # MP+H+D
    "14": {"LT": True,  "MP": True,  "H": False, "D": False},  # LT+MP
    "15": {"LT": False, "MP": True,  "H": False, "D": True},   # MP+D
    "16": {"LT": False, "MP": True,  "H": False, "D": False},  # MP
    "17": {"LT": False, "MP": True,  "H": True,  "D": False},  # MP+H
    "18": {"LT": True,  "MP": False, "H": False, "D": True},   # LT+D
    "19": {"LT": True,  "MP": True,  "H": False, "D": True},   # LT+MP+D
    "20": {"LT": True,  "MP": True,  "H": True,  "D": False},  # LT+MP+H
}

# Mô tả nội dung thi
NOIDUNG_MOTA = {
    "1":  "LT+H lần đầu A1",
    "2":  "Chỉ LT",
    "3":  "LT+H",
    "4":  "LT+H+D",
    "5":  "H+D",
    "6":  "Chỉ D",
    "7":  "Chỉ H",
    "8":  "H (miễn LT ô tô)",
    "9":  "H (miễn LT xe máy)",
    "11": "LT+MP+H+D lần đầu",
    "12": "LT+MP+H+D thi lại",
    "13": "MP+H+D",
    "14": "LT+MP",
    "15": "MP+D",
    "16": "Chỉ MP",
    "17": "MP+H",
    "18": "LT+D",
    "19": "LT+MP+D",
    "20": "LT+MP+H",
}


# ======================= LOGGING UTILITY =======================

def get_log_file(file_type):
    """Lấy đường dẫn file log dựa trên loại import."""
    mapping = {
        "Lý Thuyết": "log_import_Ly_Thuyet.txt",
        "Mô Phỏng": "log_import_Mo_Phong.txt",
        "Hình": "log_import_Hinh.txt",
        "Đường": "log_import_Duong.txt",
        "Nội Dung": "log_import_Noi_Dung.txt",
        "XML": "log_import_XML.txt",
        "General": "log_general.txt",
    }
    filename = mapping.get(file_type, "log_import_khac.txt")
    return os.path.join(LOG_DIR, filename)


def write_log(file_type, text):
    """Ghi log với timestamp vào file tương ứng."""
    try:
        log_path = get_log_file(file_type)
        ts = datetime.datetime.now().strftime("[%Y-%m-%d %H:%M:%S]")
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(f"{ts} {text.strip()}\n")
    except Exception as e:
        print(f"Lỗi ghi log: {e}")


# ======================= CONFIG LOADER =======================

class ConfigLoader:
    """Class quản lý việc load và truy xuất config"""
    
    _excel_config = None
    _noidung_config = None
    _display_config = None
    _dvhc_data = None
    _dvhc_dict = None
    
    # ==================== EXCEL CONFIG ====================
    @classmethod
    def load_excel_config(cls, file_path=None):
        """Load cấu hình Excel từ file JSON"""
        if cls._excel_config is not None and file_path is None:
            return cls._excel_config
        
        config_path = file_path or EXCEL_CONFIG_FILE
        
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                cls._excel_config = json.load(f)
            print(f"✓ Đã load config Excel từ: {config_path}")
            return cls._excel_config
        except FileNotFoundError:
            print(f"✗ Không tìm thấy file: {config_path}, sử dụng config mặc định")
            cls._excel_config = cls._get_default_excel_config()
            return cls._excel_config
        except json.JSONDecodeError as e:
            print(f"✗ Lỗi parse JSON: {e}, sử dụng config mặc định")
            cls._excel_config = cls._get_default_excel_config()
            return cls._excel_config
        except Exception as e:
            print(f"✗ Lỗi load config: {e}, sử dụng config mặc định")
            cls._excel_config = cls._get_default_excel_config()
            return cls._excel_config
    
    @classmethod
    def save_excel_config(cls, config_dict, file_path=None):
        """Lưu cấu hình Excel ra file JSON"""
        config_path = file_path or EXCEL_CONFIG_FILE
        try:
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(config_dict, f, ensure_ascii=False, indent=4)
            print(f"✓ Đã lưu config Excel vào: {config_path}")
            cls._excel_config = None
            return True
        except Exception as e:
            print(f"✗ Lỗi lưu config: {e}")
            return False
    
    # ==================== NOI DUNG CONFIG ====================
    @classmethod
    def load_noidung_config(cls, file_path=None):
        """Load cấu hình nội dung từ file JSON"""
        if cls._noidung_config is not None and file_path is None:
            return cls._noidung_config
        
        config_path = file_path or NOIDUNG_CONFIG_FILE
        
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                cls._noidung_config = json.load(f)
            print(f"✓ Đã load config nội dung từ: {config_path}")
            return cls._noidung_config
        except:
            cls._noidung_config = cls._get_default_noidung_config()
            return cls._noidung_config
    
    # ==================== DISPLAY CONFIG ====================
    @classmethod
    def load_display_config(cls, file_path=None):
        """Load cấu hình hiển thị cột"""
        if cls._display_config is not None and file_path is None:
            return cls._display_config
        
        config_path = file_path or DISPLAY_CONFIG_FILE
        
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                cls._display_config = json.load(f)
            print(f"✓ Đã load config hiển thị từ: {config_path}")
            return cls._display_config
        except:
            cls._display_config = cls._get_default_display_config()
            return cls._display_config
    
    @classmethod
    def save_display_config(cls, config_dict, file_path=None):
        """Lưu cấu hình hiển thị ra file JSON"""
        config_path = file_path or DISPLAY_CONFIG_FILE
        try:
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(config_dict, f, ensure_ascii=False, indent=4)
            cls._display_config = None
            return True
        except Exception as e:
            print(f"✗ Lỗi lưu config: {e}")
            return False
    
    # ==================== DVHC CONFIG (CFG FORMAT) ====================
    @classmethod

    def load_dvhc_data(cls, file_path=None):
        """
        Load dữ liệu đơn vị hành chính từ file CFG
        
        Định dạng file CFG:
        [DVHC]
        00004=01|0100004|Phường Ba Đình|Ba Đình|Phường Ba Đình, Thành phố Hà Nội|PH
        
        Cấu trúc value: MA_DV|MA_DVQL|TEN_DVHC|TENNGANGON|TENDAYDU|LOAIDVHC
        """
        if cls._dvhc_data is not None and file_path is None:
            return cls._dvhc_data
        
        config_path = file_path or DM_DVHC_FILE
        
        try:
            print(f"Đang load ĐVHC từ: {config_path}")
            
            # Kiểm tra file tồn tại
            if not os.path.exists(config_path):
                print(f"✗ File không tồn tại: {config_path}")
                cls._dvhc_data = []
                cls._dvhc_dict = {}
                return cls._dvhc_data
            
            cls._dvhc_data = []
            cls._dvhc_dict = {}
            
            # ĐỌC FILE THỦ CÔNG
            in_dvhc_section = False
            
            # Thử các encoding khác nhau
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            file_content = None
            
            for encoding in encodings:
                try:
                    with open(config_path, 'r', encoding=encoding) as f:
                        file_content = f.readlines()
                    print(f"✓ Đọc file thành công với encoding: {encoding}")
                    break
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    print(f"⚠️ Lỗi đọc với encoding {encoding}: {e}")
                    continue
            
            if file_content is None:
                print(f"✗ Không thể đọc file với bất kỳ encoding nào")
                return cls._dvhc_data
            
            # Parse từng dòng
            for line_num, line in enumerate(file_content, 1):
                line = line.strip()
                
                # Bỏ qua dòng trống và comment
                if not line or line.startswith(';') or line.startswith('#'):
                    continue
                
                # Kiểm tra section header
                if line.startswith('[') and line.endswith(']'):
                    section_name = line[1:-1].strip().upper()
                    in_dvhc_section = (section_name == 'DVHC')
                    continue
                
                # Chỉ xử lý nếu đang trong section DVHC
                if not in_dvhc_section:
                    continue
                
                # Parse dòng key=value
                if '=' not in line:
                    continue
                
                try:
                    # Tách key và value
                    eq_pos = line.index('=')
                    ma_dvhc = line[:eq_pos].strip()
                    value = line[eq_pos+1:].strip()
                    
                    if not ma_dvhc or not value:
                        continue
                    
                    # Tách các trường bằng dấu |
                    parts = value.split('|')
                    
                    if len(parts) >= 5:
                        item = {
                            "MA_DVHC": ma_dvhc,
                            "MA_DV": parts[0].strip() if len(parts) > 0 else "",
                            "MA_DVQL": parts[1].strip() if len(parts) > 1 else "",
                            "TEN_DVHC": parts[2].strip() if len(parts) > 2 else "",
                            "TENNGANGON": parts[3].strip() if len(parts) > 3 else "",
                            "TENDAYDU": parts[4].strip() if len(parts) > 4 else "",
                            "LOAIDVHC": parts[5].strip() if len(parts) > 5 else ""
                        }
                        
                        cls._dvhc_data.append(item)
                        cls._dvhc_dict[ma_dvhc] = item
                    
                except Exception as e:
                    print(f"⚠️ Lỗi parse dòng {line_num}: {e}")
                    continue
            
            print(f"✓ Đã load DM_DVHC: {len(cls._dvhc_data)} bản ghi")
            return cls._dvhc_data
            
        except Exception as e:
            print(f"✗ Lỗi load DM_DVHC: {e}")
            traceback.print_exc()
            cls._dvhc_data = []
            cls._dvhc_dict = {}
            return cls._dvhc_data
    
    @classmethod
    def get_dvhc_tendaydu(cls, ma_dvhc, ma_dvql=None):
        """Lấy TENDAYDU từ mã ĐVHC"""
        if cls._dvhc_data is None:
            cls.load_dvhc_data()
        
        if not cls._dvhc_dict or not ma_dvhc:
            return ""
        
        ma_dvhc = str(ma_dvhc).strip()
        
        # Thử tìm trực tiếp
        if ma_dvhc in cls._dvhc_dict:
            item = cls._dvhc_dict[ma_dvhc]
            if ma_dvql:
                ma_dvql = str(ma_dvql).strip()
                if str(item.get("MA_DVQL", "")).strip() == ma_dvql:
                    return item.get("TENDAYDU", "")
            else:
                return item.get("TENDAYDU", "")
        
        # Thử với padding số 0
        if len(ma_dvhc) < 5:
            ma_dvhc_padded = ma_dvhc.zfill(5)
            if ma_dvhc_padded in cls._dvhc_dict:
                return cls._dvhc_dict[ma_dvhc_padded].get("TENDAYDU", "")
        
        return ""
    
    @classmethod
    def get_dvhc_info(cls, ma_dvhc):
        """Lấy toàn bộ thông tin ĐVHC từ mã"""
        if cls._dvhc_data is None:
            cls.load_dvhc_data()
        
        if not cls._dvhc_dict or not ma_dvhc:
            return None
        
        ma_dvhc = str(ma_dvhc).strip()
        return cls._dvhc_dict.get(ma_dvhc, None)
    
    @classmethod
    def reload_dvhc_data(cls):
        """Reload dữ liệu ĐVHC từ file"""
        cls._dvhc_data = None
        cls._dvhc_dict = None
        return cls.load_dvhc_data()
    
    @classmethod
    def search_dvhc(cls, keyword):
        """Tìm kiếm ĐVHC theo từ khóa"""
        if cls._dvhc_data is None:
            cls.load_dvhc_data()
        
        if not cls._dvhc_data or not keyword:
            return []
        
        keyword = keyword.lower()
        results = []
        
        for item in cls._dvhc_data:
            tendaydu = str(item.get("TENDAYDU", "")).lower()
            ten_dvhc = str(item.get("TEN_DVHC", "")).lower()
            tenngangon = str(item.get("TENNGANGON", "")).lower()
            ma_dvhc = str(item.get("MA_DVHC", ""))
            
            if (keyword in tendaydu or keyword in ten_dvhc or 
                keyword in tenngangon or keyword in ma_dvhc):
                results.append(item)
        
        return results
    

    @classmethod
    def get_dvhc_by_madvql(cls, ma_dvql_full):
        """
        Tìm ĐVHC theo mã DVQL đầy đủ (NOI_CT_MA_DVQL + NOI_CT_MA_DVHC)
        
        Args:
            ma_dvql_full: Mã đầy đủ (VD: "7925207" = "79" + "25207")
        
        Returns:
            str: TENDAYDU hoặc chuỗi rỗng nếu không tìm thấy
        """
        if cls._dvhc_data is None:
            cls.load_dvhc_data()
        
        if not cls._dvhc_data or not ma_dvql_full:
            return ""
        
        ma_dvql_full = str(ma_dvql_full).strip()
        
        # Tìm trong danh sách ĐVHC theo cột MA_DVQL
        for item in cls._dvhc_data:
            item_ma_dvql = str(item.get("MA_DVQL", "")).strip()
            if item_ma_dvql == ma_dvql_full:
                return item.get("TENDAYDU", "")
        
        return ""

    @classmethod
    def get_dvhc_full_info_by_madvql(cls, ma_dvql_full):
        """
        Lấy toàn bộ thông tin ĐVHC theo mã DVQL đầy đủ
        
        Args:
            ma_dvql_full: Mã đầy đủ (VD: "7925207")
        
        Returns:
            dict hoặc None
        """
        if cls._dvhc_data is None:
            cls.load_dvhc_data()
        
        if not cls._dvhc_data or not ma_dvql_full:
            return None
        
        ma_dvql_full = str(ma_dvql_full).strip()
        
        for item in cls._dvhc_data:
            item_ma_dvql = str(item.get("MA_DVQL", "")).strip()
            if item_ma_dvql == ma_dvql_full:
                return item
        
        return None





    # ==================== RELOAD ALL ====================
    @classmethod
    def reload_configs(cls):
        """Reload tất cả config từ file"""
        cls._excel_config = None
        cls._noidung_config = None
        cls._display_config = None
        cls._dvhc_data = None
        cls._dvhc_dict = None
        cls.load_excel_config()
        cls.load_noidung_config()
        cls.load_display_config()
        cls.load_dvhc_data()
    
    # ==================== DEFAULT CONFIGS ====================
    @classmethod
    def _get_default_excel_config(cls):
        """Trả về config mặc định"""
        return {
            "Lý Thuyết": {
                "header_rows": 22,
                "columns": "B, F, G, M, Q, AB",
                "headers": "SO_TT, SO_BAO_DANH, HO_VA_TEN, SO_CMT, NGAY_SINH, HANG_GPLX",
                "sbd_column": "SO_BAO_DANH",
                "name_column": "HO_VA_TEN",
                "result_column": None,
                "db_column": "KQ_SH_LYTHUYET",
                "description": "Kết quả thi Lý thuyết"
            },
            "Mô Phỏng": {
                "header_rows": 10,
                "columns": "B, C, D, E, F, H",
                "headers": "SO_TT, SO_BAO_DANH, HO_VA_TEN, SO_CMT, NGAY_SINH, HANG_GPLX",
                "sbd_column": "SO_BAO_DANH",
                "name_column": "HO_VA_TEN",
                "result_column": None,
                "db_column": "KQ_SH_MOPHONG",
                "description": "Kết quả thi Mô phỏng"
            },
            "Hình": {
                "header_rows": 33,
                "columns": "A, D, H, W, AE, AL",
                "headers": "SO_TT, SO_BAO_DANH, HO_VA_TEN, NGAY_SINH, HANG_GPLX, TONG_DIEM",
                "sbd_column": "SO_BAO_DANH",
                "name_column": "HO_VA_TEN",
                "result_column": "TONG_DIEM",
                "db_column": "KQ_SH_HINH",
                "description": "Kết quả thi trong Hình"
            },
            "Đường": {
                "header_rows": 22,
                "columns": "A, C, H, M, S, V, Y, AA",
                "headers": "SO_TT, HANG_GPLX, SO_BAO_DANH, HO_VA_TEN, NGAY_SINH, SO_CMT, DIEM, KET_QUA",
                "sbd_column": "SO_BAO_DANH",
                "name_column": "HO_VA_TEN",
                "result_column": "KET_QUA",
                "db_column": "KQ_SH_DUONG",
                "description": "Kết quả thi trên Đường"
            },
            "Nội Dung": {
                "header_rows": 1,
                "columns": "A, B, C",
                "headers": "SO_BAO_DANH, HO_VA_TEN, NOI_DUNG",
                "sbd_column": "SO_BAO_DANH",
                "name_column": "HO_VA_TEN",
                "result_column": "NOI_DUNG",
                "db_column": "NOI_DUNG_SH",
                "description": "Nội dung sát hạch"
            }
        }
    
    @classmethod
    def _get_default_noidung_config(cls):
        """Trả về config nội dung mặc định"""
        return {"MAP_NOI_DUNG": {}, "DESCRIPTIONS": {}}
    
    @classmethod
    def _get_default_display_config(cls):
        """Trả về config hiển thị mặc định"""
        return {
            "visible_columns": ["SO_TT", "SO_BAO_DANH", "MA_DK", "HO_VA_TEN", "NGAY_SINH"],
            "column_widths": {"SO_TT": 50, "SO_BAO_DANH": 80},
            "column_labels": {"SO_TT": "STT", "SO_BAO_DANH": "SBD"},
            "show_all_columns": False,
            "auto_width": True
        }


# ======================= HELPER FUNCTIONS =======================

def get_excel_config(file_type):
    """Lấy cấu hình cho loại file Excel cụ thể"""
    configs = ConfigLoader.load_excel_config()
    return configs.get(file_type, None)

def get_all_excel_configs():
    """Lấy tất cả cấu hình Excel"""
    return ConfigLoader.load_excel_config()

def get_column_list(file_type):
    """Lấy danh sách cột dạng list từ config"""
    cfg = get_excel_config(file_type)
    if cfg:
        return [col.strip() for col in cfg.get("columns", "").split(",")]
    return []

def get_header_list(file_type):
    """Lấy danh sách header dạng list từ config"""
    cfg = get_excel_config(file_type)
    if cfg:
        return [h.strip() for h in cfg.get("headers", "").split(",")]
    return []

def column_letter_to_index(letter):
    """Chuyển đổi ký tự cột Excel sang index"""
    letter = letter.strip().upper()
    result = 0
    for char in letter:
        result = result * 26 + (ord(char) - ord('A') + 1)
    return result - 1

def get_column_indices(file_type):
    """Lấy danh sách index của các cột cần lấy từ config"""
    columns = get_column_list(file_type)
    return [column_letter_to_index(col) for col in columns]

def get_noi_dung_map():
    """Lấy map nội dung sát hạch"""
    cfg = ConfigLoader.load_noidung_config()
    return cfg.get("MAP_NOI_DUNG", {})

def get_display_config():
    """Lấy cấu hình hiển thị"""
    return ConfigLoader.load_display_config()

def get_dvhc_data():
    """Lấy toàn bộ dữ liệu ĐVHC"""
    return ConfigLoader.load_dvhc_data()

def get_tendaydu_by_ma(ma_dvhc, ma_dvql=None):
    """Lấy TENDAYDU từ mã ĐVHC và mã ĐVQL"""
    return ConfigLoader.get_dvhc_tendaydu(ma_dvhc, ma_dvql)

def get_dvhc_info(ma_dvhc):
    """Lấy thông tin đầy đủ của ĐVHC theo mã"""
    return ConfigLoader.get_dvhc_info(ma_dvhc)

def search_dvhc(keyword):
    """Tìm kiếm ĐVHC theo từ khóa"""
    return ConfigLoader.search_dvhc(keyword)


# ======================= MAIN APPLICATION =======================
class ExamDataManager:
    def __init__(self, root):
        self.root = root
        self.root.title("Trường TCN Giao Thông và Xây Dựng Việt Úc")
        self.root.geometry("1100x800")

        # Database
        self.db_path = DB_FILE
        self.conn = None
        self.current_columns = []

        # Khởi tạo file_configs
        self.file_configs = ConfigLoader.load_excel_config()
        
        # Load DVHC
        print("\n" + "="*50)
        print("KHỞI ĐỘNG - LOAD DỮ LIỆU ĐVHC")
        print("="*50)
        dvhc_data = ConfigLoader.load_dvhc_data()
        print(f"Số bản ghi ĐVHC: {len(dvhc_data) if dvhc_data else 0}")
        print("="*50 + "\n")

        # Kết nối và tạo bảng
        self.create_connection()
        self.create_tables()

        # Build UI
        self.create_widgets()

        # Load initial data
        self.refresh_data()
    
    # NOTE: Đã chuyển hàm get_log_file() và write_log() ra ngoài lớp
    # để chúng có thể được dùng dễ dàng và không bị ràng buộc bởi instance của class.


    # ---------------------------
    # Database helpers
    # ---------------------------
    def create_connection(self):
        try:
            self.conn = sqlite3.connect(self.db_path)
            write_log("General", "Kết nối CSDL thành công.")
        except Exception as e:
            messagebox.showerror("Lỗi DB", f"Không thể kết nối tới DB: {e}")
            write_log("General", f"LỖI: Không thể kết nối CSDL: {e}")
            self.conn = None

    def get_connection(self):
        if not self.conn:
            self.create_connection()
        return self.conn

    def create_tables(self):
        conn = self.get_connection()
        if not conn:
            return
        c = conn.cursor()
        c.execute("""
            CREATE TABLE IF NOT EXISTS NguoiLX_HoSo (
                SO_TT INTEGER, MA_DK TEXT, HO_TEN_DEM TEXT, TEN TEXT, HO_VA_TEN TEXT,
                GIOI_TINH TEXT, NGAY_SINH DATE, MA_QUOC_TICH TEXT, NOI_CT TEXT, NOI_CT_MA_DVHC TEXT,
                NOI_CT_MA_DVQL TEXT, SO_CMT TEXT, SO_HO_SO TEXT, MA_KY_SH TEXT, SO_BAO_DANH INTEGER,
                MA_CSDT TEXT, MA_TTSH TEXT, MA_SO_GTVT TEXT, GIAY_CNSK TEXT, HANG_GPLX TEXT,
                SO_GPLX_DA_CO TEXT, HANG_GPLX_DA_CO TEXT, DVQL_GPLX_DACO TEXT, NGAY_HH_GPLX_DACO DATE,
                SO_NAM_LAIXE INTEGER, SO_KM_ANTOAN INTEGER, SO_GIAY_CNTN TEXT, SO_CCN TEXT, NOI_DUNG_SH TEXT,
                LY_DO_SH TEXT, KET_QUA_SH TEXT, KQ_SH_LYTHUYET TEXT, KQ_SH_MOPHONG TEXT, KQ_SH_HINH TEXT,
                KQ_SH_DUONG TEXT, GHI_CHU_SH TEXT, ANH_CHAN_DUNG TEXT, NGAY_TT_GPLX_DACO TEXT,
                MA_KHOA_HOC TEXT, SO_QD_SH TEXT, NGAY_QD_SH DATE, NGUOI_QD_SH TEXT, CHAT_LUONG_ANH TEXT,
                LYTHUYETKT TEXT, MOPHONGKT TEXT, HINHKT TEXT, DUONGKT TEXT, KETQUAKT TEXT, TRANGTHAI TEXT
            )
        """)
        c.execute("""
            CREATE TABLE IF NOT EXISTS KY_SH (
                MAKYSH TEXT, MATTSH TEXT, NGAYSH DATE, GIOSH TEXT, SOQD TEXT,
                NGAYQD DATE, NGUOIQD TEXT, CHUTICH_HDSH TEXT, PHOCHUTICH_HDSH TEXT,
                UV_GD_TTSH TEXT, UV_TOTRUONG TEXT, UV_THUKY TEXT, TONGSODK TEXT
            )
        """)
        conn.commit()
        write_log("General", "Kiểm tra/Tạo các bảng NguoiLX_HoSo và KY_SH thành công.")

    #Thêm hàm clear_sbd_input

    def clear_sbd_input(self):
        """Xóa ô nhập SBD và label tên"""
        self.entry_input_sbd.delete(0, tk.END)
        self.lbl_student_name.config(text="")
        self.entry_input_sbd.focus_set()

    # Thêm hàm show_noidung_map
    def show_noidung_map(self):
        """Hiển thị bảng nội dung thi"""
        noidung_win = Toplevel(self.root)
        noidung_win.title("Bảng Nội dung Thi")
        noidung_win.geometry("700x500")
        noidung_win.resizable(True, True)
        
        Label(noidung_win, text="BẢNG NỘI DUNG THI SÁT HẠCH",
            font=("Arial", 14, "bold"), fg="blue").pack(pady=10)
        
        # Frame chứa bảng
        tree_frame = Frame(noidung_win)
        tree_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)
        
        # Scrollbar
        scroll_y = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
        
        # Treeview
        columns = ("ma", "mota", "lt", "mp", "h", "d")
        tree = ttk.Treeview(tree_frame, columns=columns, show="headings",
                        yscrollcommand=scroll_y.set, height=18)
        
        scroll_y.config(command=tree.yview)
        
        # Cấu hình cột
        tree.heading("ma", text="Mã ND")
        tree.heading("mota", text="Mô tả")
        tree.heading("lt", text="Lý thuyết")
        tree.heading("mp", text="Mô phỏng")
        tree.heading("h", text="Hình")
        tree.heading("d", text="Đường")
        
        tree.column("ma", width=60, anchor="center")
        tree.column("mota", width=250, anchor="w")
        tree.column("lt", width=80, anchor="center")
        tree.column("mp", width=80, anchor="center")
        tree.column("h", width=80, anchor="center")
        tree.column("d", width=80, anchor="center")
        
        # Thêm dữ liệu
        for ma, req in sorted(NOIDUNG_MAP.items(), key=lambda x: int(x[0]) if x[0].isdigit() else 999):
            mota = NOIDUNG_MOTA.get(ma, "")
            lt = "✅ Thi" if req.get("LT") else "❌ Không"
            mp = "✅ Thi" if req.get("MP") else "❌ Không"
            h = "✅ Thi" if req.get("H") else "❌ Không"
            d = "✅ Thi" if req.get("D") else "❌ Không"
            
            tree.insert("", "end", values=(ma, mota, lt, mp, h, d))
        
        tree.pack(side=LEFT, fill=BOTH, expand=True)
        scroll_y.pack(side=RIGHT, fill=Y)
        
        # Chú thích
        Label(noidung_win, text="✅ Thi = Phải thi môn đó | ❌ Không = Không thi môn đó",
            font=("Arial", 10, "italic"), fg="gray").pack(pady=5)
        
        Button(noidung_win, text="Đóng", command=noidung_win.destroy, width=15).pack(pady=10)


    # ---------------------------
    # UI
    # ---------------------------
    def create_widgets(self):
        """Tạo giao diện người dùng"""
        
        # ========== CẤU HÌNH ROOT WINDOW ==========
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        
        # ==================== MAIN FRAME ====================
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Cấu hình main_frame
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(4, weight=1)  # Data frame sẽ giãn

        # ==================== ROW 0: TITLE ====================
        title_label = ttk.Label(
            main_frame, 
            text="Phần mềm Quản lý Dữ liệu thi Sát Hạch Thiên Tôn",
            font=("Arial", 16, "bold"),
            foreground="blue",
            wraplength=800,  # ⭐ THÊM WRAPLENGTH
            anchor="center"
        )
        title_label.grid(row=0, column=0, pady=10, sticky=(tk.W, tk.E))

        # ==================== ROW 1: Nhập Xuất File ====================
        func_frame = ttk.Frame(main_frame)
        func_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=5)
        
        # Cấu hình func_frame để các cột có weight bằng nhau
        func_frame.columnconfigure(0, weight=1)
        func_frame.columnconfigure(1, weight=1)
        func_frame.columnconfigure(2, weight=1)

        # ----- col=0: Nhập dữ liệu đầu vào -----
        import_frame = ttk.LabelFrame(func_frame, text="📥 Nhập Xuất File DL đầu vào", padding="10")
        import_frame.grid(row=0, column=0, padx=5, pady=5, sticky=(tk.N, tk.S, tk.W, tk.E))

        ttk.Button(import_frame, text="Nhập File thi PC08", 
                command=self.import_xml, width=20).grid(row=0, column=0, padx=3, pady=3)
        ttk.Button(import_frame, text="📊 Xuất DL Excel", 
                command=self.export_to_excel, width=20).grid(row=1, column=0, padx=3, pady=3)
        ttk.Button(import_frame, text="🗑️ Xóa dữ liệu", 
                command=self.clear_data, width=20).grid(row=0, column=1, padx=3, pady=3)

        # ----- col=1: Thống kê kết quả nhập -----
        export_frame = ttk.LabelFrame(func_frame, text="📤 Thống kê kết quả nhập", padding="10")
        export_frame.grid(row=0, column=1, padx=5, pady=5, sticky=(tk.N, tk.S, tk.W, tk.E))

        ttk.Button(export_frame, text="📊 Xem kết quả nhập", 
                command=self.view_import_results, width=20).grid(row=0, column=0, padx=3, pady=3)
        ttk.Button(export_frame, text="📋 DS chưa có KQ", 
                command=self.view_missing_results, width=20).grid(row=1, column=0, padx=3, pady=3)
        ttk.Button(export_frame, text="📄 Xuất báo cáo KQ", 
                command=self.export_result_report, width=20).grid(row=0, column=1, padx=3, pady=3)

        # ----- col=2: config_frame -----
        config_frame = ttk.LabelFrame(func_frame, text="⚙️ Cấu Hình", padding="10")
        config_frame.grid(row=0, column=2, padx=5, pady=5, sticky=(tk.N, tk.S, tk.W, tk.E))

        ttk.Button(config_frame, text="📍 Quản lý ĐVHC", 
                command=self.manage_dvhc, width=20).grid(row=0, column=0, padx=3, pady=3)
        ttk.Button(config_frame, text="🔄 Cập nhật ĐVHC", 
                command=self.update_all_noict_from_dvhc, width=20).grid(row=1, column=0, padx=3, pady=3)
        ttk.Button(config_frame, text="📅 Xem Kỳ SH", 
                command=self.view_kysh, width=20).grid(row=0, column=1, padx=3, pady=3)
        ttk.Button(config_frame, text="⚙️ Cấu hình Excel", 
                command=self.open_config_editor, width=20).grid(row=1, column=1, padx=3, pady=3)

        # ==================== ROW 2: SATHACH_FRAME ====================
        Ketqua_frame = ttk.LabelFrame(main_frame, text="Nhập kết quả thi sát hạch của thí sinh", padding="10")
        Ketqua_frame.grid(row=3, column=0, pady=10, sticky=(tk.W, tk.E))
        
        # Cấu hình Ketqua_frame
        for i in range(6):
            Ketqua_frame.columnconfigure(i, weight=1)

        # --- NHẬP SBD ---
        search_frame = ttk.Frame(Ketqua_frame)
        search_frame.grid(row=0, column=0, padx=5, pady=5, sticky=(tk.N, tk.S))

        ttk.Label(search_frame, text="Nhập SBD:", font=("Arial", 12, "bold"), foreground="red") \
            .grid(row=0, column=0, padx=5, pady=2, sticky="w")

        self.entry_input_sbd = ttk.Entry(search_frame, width=20, font=("Arial", 14))
        self.entry_input_sbd.grid(row=1, column=0, padx=5, pady=5)

        # Nút xóa SBD
        ttk.Button(search_frame, text="🗑️", width=3, 
                command=self.clear_sbd_input).grid(row=1, column=1, padx=2, pady=5)

        # Label hiển thị tên học viên
        self.lbl_student_name = ttk.Label(search_frame, text="", font=("Arial", 9), 
                                        foreground="blue", wraplength=180)
        self.lbl_student_name.grid(row=2, column=0, columnspan=2, padx=5, pady=2, sticky="w")

        # Bind sự kiện
        self.entry_input_sbd.bind("<KeyRelease>", self.on_sbd_input_change)
        self.entry_input_sbd.bind("<Return>", lambda e: self.focus_next_input())
        self.entry_input_sbd.bind("<Escape>", lambda e: self.clear_sbd_input())

        # Lý Thuyết
        lt_frame = ttk.Frame(Ketqua_frame)
        lt_frame.grid(row=0, column=1, padx=5, pady=5, sticky=(tk.N, tk.S))
        ttk.Label(lt_frame, text="Lý Thuyết:", font=("Arial", 10, "bold")).pack()
        ttk.Button(lt_frame, text="LT Đạt", command=lambda: self.update_result_by_sbd("LYTHUYETKT", "DA"), width=15).pack(pady=2)
        ttk.Button(lt_frame, text="LT Trượt", command=lambda: self.update_result_by_sbd("LYTHUYETKT", "RO"), width=15).pack(pady=2)
        ttk.Button(lt_frame, text="LT Vắng", command=lambda: self.update_result_by_sbd("LYTHUYETKT", "VA"), width=15).pack(pady=2)

        # Mô Phỏng
        mp_frame = ttk.Frame(Ketqua_frame)
        mp_frame.grid(row=0, column=2, padx=5, pady=5, sticky=(tk.N, tk.S))
        ttk.Label(mp_frame, text="Mô Phỏng:", font=("Arial", 10, "bold")).pack()
        ttk.Button(mp_frame, text="MP Đạt", command=lambda: self.update_result_by_sbd("MOPHONGKT", "DA"), width=15).pack(pady=2)
        ttk.Button(mp_frame, text="MP Trượt", command=lambda: self.update_result_by_sbd("MOPHONGKT", "RO"), width=15).pack(pady=2)
        ttk.Button(mp_frame, text="MP Vắng", command=lambda: self.update_result_by_sbd("MOPHONGKT", "VA"), width=15).pack(pady=2)

        # Hình
        h_frame = ttk.Frame(Ketqua_frame)
        h_frame.grid(row=0, column=3, padx=5, pady=5, sticky=(tk.N, tk.S))
        ttk.Label(h_frame, text="Hình:", font=("Arial", 10, "bold")).pack()
        ttk.Button(h_frame, text="H Đạt", command=lambda: self.update_result_by_sbd("HINHKT", "DA"), width=15).pack(pady=2)
        ttk.Button(h_frame, text="H Trượt", command=lambda: self.update_result_by_sbd("HINHKT", "RO"), width=15).pack(pady=2)
        ttk.Button(h_frame, text="H Vắng", command=lambda: self.update_result_by_sbd("HINHKT", "VA"), width=15).pack(pady=2)

        # Đường
        d_frame = ttk.Frame(Ketqua_frame)
        d_frame.grid(row=0, column=4, padx=5, pady=5, sticky=(tk.N, tk.S))
        ttk.Label(d_frame, text="Đường:", font=("Arial", 10, "bold")).pack()
        ttk.Button(d_frame, text="Đ Đạt", command=lambda: self.update_result_by_sbd("DUONGKT", "DA"), width=15).pack(pady=2)
        ttk.Button(d_frame, text="Đ Trượt", command=lambda: self.update_result_by_sbd("DUONGKT", "RO"), width=15).pack(pady=2)
        ttk.Button(d_frame, text="Đ Vắng", command=lambda: self.update_result_by_sbd("DUONGKT", "VA"), width=15).pack(pady=2)

        # --- Tìm kiếm nhanh ---
        search_frame2 = ttk.Frame(Ketqua_frame)
        search_frame2.grid(row=0, column=5, padx=5, pady=5, sticky=(tk.N, tk.S))

        ttk.Label(search_frame2, text="Tìm kiếm:", font=("Arial", 10, "bold")) \
            .grid(row=0, column=0, columnspan=2, padx=5, pady=2, sticky="w")

        self.entry_search_sbd = ttk.Entry(search_frame2, width=15)
        self.entry_search_sbd.grid(row=1, column=0, padx=2, pady=2)
        ttk.Button(search_frame2, text="🔍 SBD", width=8, command=self.search_sbd) \
            .grid(row=1, column=1, padx=2, pady=2)

        self.entry_search_name = ttk.Entry(search_frame2, width=15)
        self.entry_search_name.grid(row=2, column=0, padx=2, pady=2)
        ttk.Button(search_frame2, text="🔍 Tên", width=8, command=self.search_ho_ten) \
            .grid(row=2, column=1, padx=2, pady=2)

        # Nút xem bảng nội dung
        ttk.Button(search_frame2, text="📋 Bảng ND", width=25, command=self.show_noidung_map) \
            .grid(row=3, column=0, columnspan=2, padx=2, pady=2)

        # Bind phím Enter cho tìm kiếm
        self.entry_search_sbd.bind("<Return>", lambda e: self.search_sbd())
        self.entry_search_name.bind("<Return>", lambda e: self.search_ho_ten())

        # ==================== ROW 4: DATA DISPLAY ====================
        data_frame = ttk.LabelFrame(main_frame, text="Dữ liệu NguoiLX_HoSo", padding="10")
        data_frame.grid(row=4, column=0, pady=10, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Cấu hình data_frame
        data_frame.columnconfigure(0, weight=1)
        data_frame.rowconfigure(0, weight=1)

        # Treeview for data
        self.tree_scroll_y = ttk.Scrollbar(data_frame, orient=tk.VERTICAL)
        self.tree_scroll_x = ttk.Scrollbar(data_frame, orient=tk.HORIZONTAL)

        self.data_tree = ttk.Treeview(data_frame,
                                    yscrollcommand=self.tree_scroll_y.set,
                                    xscrollcommand=self.tree_scroll_x.set,
                                    height=12)
        
        # Thêm sự kiện Double Click để mở cửa sổ chỉnh sửa
        self.data_tree.bind("<Double-1>", self.edit_item)

        self.tree_scroll_y.config(command=self.data_tree.yview)
        self.tree_scroll_x.config(command=self.data_tree.xview)

        self.data_tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.tree_scroll_y.grid(row=0, column=1, sticky=(tk.N, tk.S))
        self.tree_scroll_x.grid(row=1, column=0, sticky=(tk.W, tk.E))

        # Refresh button
        ttk.Button(data_frame, text="🔄 Load lại dữ liệu", command=self.refresh_data).grid(row=2, column=0, pady=5)

        # ==================== ROW 5: STATUS BAR ====================
        self.status_label = ttk.Label(main_frame, text="Sẵn sàng", relief=tk.SUNKEN, anchor=tk.W)
        self.status_label.grid(row=5, column=0, sticky=(tk.W, tk.E), pady=5)


    def export_kq_baocao(self):
            """Xuất file kết quả Báo cáo"""
            messagebox.showinfo("Thông báo", "Chức năng xuất KQ Báo cáo đang phát triển!")




    # ---------------------------
    # config (for LT/MP/H/Đ)
    # ---------------------------
    def open_config_editor(self):
        """Mở cửa sổ chỉnh sửa cấu hình"""
        config_win = Toplevel(self.root)
        config_win.title("Cấu hình đọc file Excel")
        config_win.geometry("900x650")
        config_win.resizable(True, True)

        Label(config_win, text="CẤU HÌNH ĐỌC FILE EXCEL",
              font=("Arial", 14, "bold"), fg="blue").pack(pady=10)

        main_frame = Frame(config_win)
        main_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)

        canvas = Canvas(main_frame)
        scrollbar = Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        scroll_frame = Frame(canvas)

        scroll_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scroll_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        scrollbar.pack(side=RIGHT, fill=Y)
        canvas.pack(side=LEFT, fill=BOTH, expand=True)

        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        canvas.bind_all("<MouseWheel>", on_mousewheel)

        entries = {}
        row_idx = 0
        
        for file_type, cfg in self.file_configs.items():
            Label(scroll_frame, text=f"━━━ {file_type} ━━━",
                  font=("Arial", 11, "bold"), fg="green").grid(row=row_idx, column=0, columnspan=4, pady=10, sticky=W)
            row_idx += 1

            entries[file_type] = {}

            Label(scroll_frame, text="Số dòng header:", width=18, anchor=W).grid(row=row_idx, column=0, padx=5, pady=2, sticky=W)
            e_header = Entry(scroll_frame, width=10)
            e_header.insert(0, str(cfg.get("header_rows", 0)))
            e_header.grid(row=row_idx, column=1, padx=5, pady=2, sticky=W)
            entries[file_type]["header_rows"] = e_header

            Label(scroll_frame, text="Cột (A, B, C...):", width=18, anchor=W).grid(row=row_idx, column=2, padx=5, pady=2, sticky=W)
            e_cols = Entry(scroll_frame, width=40)
            e_cols.insert(0, cfg.get("columns", ""))
            e_cols.grid(row=row_idx, column=3, padx=5, pady=2, sticky=W)
            entries[file_type]["columns"] = e_cols
            row_idx += 1

            Label(scroll_frame, text="Tên cột:", width=18, anchor=W).grid(row=row_idx, column=0, padx=5, pady=2, sticky=W)
            e_headers = Entry(scroll_frame, width=70)
            e_headers.insert(0, cfg.get("headers", ""))
            e_headers.grid(row=row_idx, column=1, columnspan=3, padx=5, pady=2, sticky=W)
            entries[file_type]["headers"] = e_headers
            row_idx += 1

            Label(scroll_frame, text="Cột SBD:", width=18, anchor=W).grid(row=row_idx, column=0, padx=5, pady=2, sticky=W)
            e_sbd = Entry(scroll_frame, width=20)
            e_sbd.insert(0, cfg.get("sbd_column", "SO_BAO_DANH"))
            e_sbd.grid(row=row_idx, column=1, padx=5, pady=2, sticky=W)
            entries[file_type]["sbd_column"] = e_sbd

            Label(scroll_frame, text="Cột Họ tên:", width=18, anchor=W).grid(row=row_idx, column=2, padx=5, pady=2, sticky=W)
            e_name = Entry(scroll_frame, width=20)
            e_name.insert(0, cfg.get("name_column", "HO_VA_TEN"))
            e_name.grid(row=row_idx, column=3, padx=5, pady=2, sticky=W)
            entries[file_type]["name_column"] = e_name
            row_idx += 1

            Label(scroll_frame, text="Cột DB cập nhật:", width=18, anchor=W).grid(row=row_idx, column=0, padx=5, pady=2, sticky=W)
            e_db = Entry(scroll_frame, width=20)
            e_db.insert(0, cfg.get("db_column", ""))
            e_db.grid(row=row_idx, column=1, padx=5, pady=2, sticky=W)
            entries[file_type]["db_column"] = e_db
            
            Label(scroll_frame, text="Cột kết quả:", width=18, anchor=W).grid(row=row_idx, column=2, padx=5, pady=2, sticky=W)
            e_result = Entry(scroll_frame, width=20)
            e_result.insert(0, cfg.get("result_column", "") or "")
            e_result.grid(row=row_idx, column=3, padx=5, pady=2, sticky=W)
            entries[file_type]["result_column"] = e_result
            row_idx += 1

        btn_frame = Frame(config_win)
        btn_frame.pack(fill=X, pady=10)

        def save_config():
            new_config = {}
            for file_type, entry_dict in entries.items():
                result_col = entry_dict["result_column"].get().strip()
                new_config[file_type] = {
                    "header_rows": int(entry_dict["header_rows"].get() or 0),
                    "columns": entry_dict["columns"].get().strip(),
                    "headers": entry_dict["headers"].get().strip(),
                    "sbd_column": entry_dict["sbd_column"].get().strip(),
                    "name_column": entry_dict["name_column"].get().strip(),
                    "db_column": entry_dict["db_column"].get().strip(),
                    "result_column": result_col if result_col else None,
                    "description": self.file_configs.get(file_type, {}).get("description", "")
                }

            if ConfigLoader.save_excel_config(new_config):
                self.file_configs = new_config
                messagebox.showinfo("Thành công", f"Đã lưu cấu hình!")
                write_log("General", f"Đã lưu cấu hình Excel")
                canvas.unbind_all("<MouseWheel>")
                config_win.destroy()
            else:
                messagebox.showerror("Lỗi", "Không thể lưu file!")

        def on_close():
            canvas.unbind_all("<MouseWheel>")
            config_win.destroy()

        config_win.protocol("WM_DELETE_WINDOW", on_close)

        Button(btn_frame, text="💾 Lưu cấu hình", command=save_config, bg="#90EE90", width=20, height=2).pack(side=LEFT, padx=20)
        Button(btn_frame, text="🔄 Reload", command=lambda: [ConfigLoader.reload_configs(), messagebox.showinfo("OK", "Đã reload!")], bg="#87CEEB", width=20, height=2).pack(side=LEFT, padx=20)
        Button(btn_frame, text="❌ Hủy", command=on_close, bg="#FFB6C1", width=20, height=2).pack(side=RIGHT, padx=20)

    # ---------------------------
    # ĐỌC EXCEL THEO CONFIG
    # ---------------------------
    def read_excel_with_config(self, file_path, file_type):
        """Đọc file Excel theo cấu hình từ config"""
        cfg = self.file_configs.get(file_type)
        if not cfg:
            raise ValueError(f"Không tìm thấy cấu hình cho loại file: {file_type}")

        write_log(file_type, f"Đọc file với config: header_rows={cfg.get('header_rows')}, columns={cfg.get('columns')}")

        header_rows = cfg.get("header_rows", 0)
        columns_str = cfg.get("columns", "")
        headers_str = cfg.get("headers", "")

        column_letters = [col.strip() for col in columns_str.split(",") if col.strip()]
        column_indices = [column_letter_to_index(col) for col in column_letters]
        headers = [h.strip() for h in headers_str.split(",") if h.strip()]

        try:
            df_raw = pd.read_excel(file_path, header=None)
            write_log(file_type, f"Đọc được {len(df_raw)} dòng, {len(df_raw.columns)} cột")

            if header_rows > 0:
                df_raw = df_raw.iloc[header_rows:]

            df_raw = df_raw.reset_index(drop=True)

            selected_columns = []
            for idx in column_indices:
                if idx < len(df_raw.columns):
                    selected_columns.append(df_raw.iloc[:, idx])
                else:
                    selected_columns.append(pd.Series([None] * len(df_raw)))

            df_result = pd.concat(selected_columns, axis=1)

            if len(headers) == len(df_result.columns):
                df_result.columns = headers
            else:
                df_result.columns = headers[:len(df_result.columns)] if len(headers) > len(df_result.columns) else headers + [f"Col_{i}" for i in range(len(headers), len(df_result.columns))]

            df_result = df_result.dropna(how='all')

            return df_result

        except Exception as e:
            write_log(file_type, f"Lỗi đọc file: {e}")
            raise


    # ---------------------------
    # Import XML - TRA CỨU NOI_CT TỪ DM_DVHC
    # ---------------------------
    def import_xml(self):
        """Import file XML và tự động tra cứu NOI_CT từ DM_DVHC.cfg"""
        file_path = filedialog.askopenfilename(filetypes=[("XML files", "*.xml")])
        if not file_path:
            return

        write_log("XML", f"Bắt đầu import file XML: {os.path.basename(file_path)}")
        self.status_label.config(text="Đang import XML...")
        self.root.update()

        try:
            tree_xml = ET.parse(file_path)
            root = tree_xml.getroot()

            conn = self.get_connection()
            cur = conn.cursor()

            # ========== LOAD DỮ LIỆU ĐVHC ==========
            dvhc_data = ConfigLoader.load_dvhc_data()
            dvhc_loaded = len(dvhc_data) > 0
            if dvhc_loaded:
                write_log("XML", f"Đã load {len(dvhc_data)} bản ghi từ DM_DVHC.cfg")
            else:
                write_log("XML", "CẢNH BÁO: Không tìm thấy hoặc không load được DM_DVHC.cfg")

            # ========== KY_SH ==========
            kysh_node = root.find(".//KY_SH")
            if kysh_node is not None:
                cur.execute("DELETE FROM KY_SH")
                def get_kysh(tag):
                    el = kysh_node.find(tag)
                    return el.text.strip() if el is not None and el.text else ""
                data_kysh = [
                    get_kysh("MAKYSH"), get_kysh("MATTSH"), get_kysh("NGAYSH"), get_kysh("GIOSH"),
                    get_kysh("SOQD"), get_kysh("NGAYQD"), get_kysh("NGUOIQD"), get_kysh("CHUTICH_HDSH"),
                    get_kysh("PHOCHUTICH_HDSH"), get_kysh("UV_GD_TTSH"), get_kysh("UV_TOTRUONG"),
                    get_kysh("UV_THUKY"), get_kysh("TONGSODK")
                ]
                cur.execute("INSERT INTO KY_SH VALUES (" + ",".join(["?"]*13) + ")", data_kysh)
                write_log("XML", f"Đã import KY_SH: {data_kysh[0]} - {data_kysh[2]}")

            # ========== NguoiLX_HoSo ==========
            nguoi_nodes = root.findall(".//NGUOI_LX")
            if not nguoi_nodes:
                messagebox.showerror("Lỗi", "Không tìm thấy thẻ <NGUOI_LX> trong file XML.")
                return

            cur.execute("DELETE FROM NguoiLX_HoSo")
            inserted = 0
            dvhc_updated = 0
            dvhc_not_found = 0
            dvhc_log_details = []

            def get_text(node, tag):
                el = node.find(tag)
                return el.text.strip() if el is not None and el.text else ""

            for nguoi in nguoi_nodes:
                hs = nguoi.find("HO_SO")
                def get(tag): return get_text(nguoi, tag)
                def get_hoso(tag): return get_text(hs, tag) if hs is not None else ""

                # Lấy thông tin cơ bản
                so_bao_danh = get_hoso("SO_BAO_DANH")
                ho_va_ten = get("HO_VA_TEN")
                
                # ========== LẤY MÃ ĐVHC TỪ XML ==========
                ma_dvql = get("NOI_CT_MA_DVQL")  # VD: "79"
                ma_dvhc = get("NOI_CT_MA_DVHC")  # VD: "25207"
                noi_ct_xml = get("NOI_CT")       # Giá trị gốc từ XML
                
                # ========== GHÉP MÃ VÀ TRA CỨU ==========
                noi_ct_final = noi_ct_xml  # Mặc định dùng giá trị từ XML
                
                if dvhc_loaded and ma_dvql and ma_dvhc:
                    # Ghép mã: NOI_CT_MA_DVQL + NOI_CT_MA_DVHC
                    ma_dvql_full = str(ma_dvql).strip() + str(ma_dvhc).strip()
                    
                    # Tìm theo mã DVQL đầy đủ (cột thứ 3 trong file cfg)
                    tendaydu = ConfigLoader.get_dvhc_by_madvql(ma_dvql_full)
                    
                    if tendaydu:
                        noi_ct_final = tendaydu
                        dvhc_updated += 1
                        dvhc_log_details.append(
                            f"✓ SBD:{so_bao_danh} | {ho_va_ten} | "
                            f"Mã:{ma_dvql}+{ma_dvhc}={ma_dvql_full} → {tendaydu[:50]}..."
                        )
                    else:
                        dvhc_not_found += 1
                        dvhc_log_details.append(
                            f"✗ SBD:{so_bao_danh} | {ho_va_ten} | "
                            f"Mã:{ma_dvql}+{ma_dvhc}={ma_dvql_full} không tìm thấy"
                        )

                # Tạo dữ liệu cho INSERT
                data = [
                    get("SO_TT"), get("MA_DK"), get("HO_TEN_DEM"), get("TEN"), ho_va_ten,
                    get("GIOI_TINH"), get("NGAY_SINH"), get("MA_QUOC_TICH"), 
                    noi_ct_final,  # NOI_CT đã tra cứu
                    ma_dvhc,       # NOI_CT_MA_DVHC
                    ma_dvql,       # NOI_CT_MA_DVQL
                    get("SO_CMT"), get_hoso("SO_HO_SO"), get_hoso("MA_KY_SH"),
                    so_bao_danh, get_hoso("MA_CSDT"), get_hoso("MA_TTSH"), get_hoso("MA_SO_GTVT"),
                    get_hoso("GIAY_CNSK"), get_hoso("HANG_GPLX"), get_hoso("SO_GPLX_DA_CO"), get_hoso("HANG_GPLX_DA_CO"),
                    get_hoso("DVQL_GPLX_DACO"), get_hoso("NGAY_HH_GPLX_DACO"), get_hoso("SO_NAM_LAIXE"),
                    get_hoso("SO_KM_ANTOAN"), get_hoso("SO_GIAY_CNTN"), get_hoso("SO_CCN"), get_hoso("NOI_DUNG_SH"),
                    get_hoso("LY_DO_SH"), get_hoso("KET_QUA_SH"), get_hoso("KQ_SH_LYTHUYET"), get_hoso("KQ_SH_MOPHONG"),
                    get_hoso("KQ_SH_HINH"), get_hoso("KQ_SH_DUONG"), get_hoso("GHI_CHU_SH"), get_hoso("ANH_CHAN_DUNG"),
                    get_hoso("NGAY_TT_GPLX_DACO"), get_hoso("MA_KHOA_HOC"), get_hoso("SO_QD_SH"), get_hoso("NGAY_QD_SH"),
                    get_hoso("NGUOI_QD_SH"), get_hoso("CHAT_LUONG_ANH"),
                    get_hoso("LYTHUYETKT"), get_hoso("MOPHONGKT"), get_hoso("HINHKT"), get_hoso("DUONGKT"), get_hoso("KETQUAKT"), get_hoso("TRANGTHAI")
                ]

                if len(data) == 49:
                    cur.execute("INSERT INTO NguoiLX_HoSo VALUES (" + ",".join(["?"]*49) + ")", data)
                    inserted += 1

            conn.commit()
            
            # ========== GHI LOG ==========
            write_log("XML", f"{'='*60}")
            write_log("XML", f"HOÀN TẤT IMPORT XML")
            write_log("XML", f"File: {os.path.basename(file_path)}")
            write_log("XML", f"Đã nhập: {inserted} hồ sơ")
            write_log("XML", f"{'='*60}")
            write_log("XML", f"TRA CỨU ĐỊA CHỈ (MA_DVQL + MA_DVHC → TENDAYDU):")
            write_log("XML", f"  - Thành công: {dvhc_updated}")
            write_log("XML", f"  - Không tìm thấy: {dvhc_not_found}")
            write_log("XML", f"{'='*60}")
            
            # Ghi chi tiết (giới hạn 50 dòng)
            for detail in dvhc_log_details[:50]:
                write_log("XML", detail)
            if len(dvhc_log_details) > 50:
                write_log("XML", f"... và {len(dvhc_log_details) - 50} bản ghi khác")

            # ========== THÔNG BÁO ==========
            msg = f"✅ Đã nhập {inserted} học viên vào CSDL.\n\n"
            
            if dvhc_loaded:
                msg += f"📍 Tra cứu địa chỉ (DVQL+DVHC → TENDAYDU):\n"
                msg += f"   • Thành công: {dvhc_updated}\n"
                msg += f"   • Không tìm thấy: {dvhc_not_found}\n"
                
                if dvhc_not_found > 0:
                    msg += f"\n⚠️ Có {dvhc_not_found} mã không tìm thấy.\n"
                    msg += "   Xem chi tiết trong file log."
            else:
                msg += "⚠️ Không load được file DM_DVHC.cfg\n"
                msg += "   Địa chỉ sẽ giữ nguyên từ file XML."
            
            messagebox.showinfo("Thành công", msg)
            self.status_label.config(text=f"Import XML: {inserted} hồ sơ | ĐVHC: {dvhc_updated}/{dvhc_updated + dvhc_not_found}")
            self.refresh_data()

        except Exception as e:
            error_msg = f"LỖI khi import XML: {e}\n\n{traceback.format_exc()}"
            messagebox.showerror("Lỗi", error_msg)
            write_log("XML", error_msg)
            self.status_label.config(text="Lỗi import XML")

    # ---------------------------
    # Refresh data
    # ---------------------------
    def refresh_data(self):
        """Làm mới dữ liệu với độ rộng cột tự động theo nội dung"""
        conn = self.get_connection()
        if not conn:
            return
        try:
            cur = conn.cursor()

            # Lấy dữ liệu
            cur.execute("""
                SELECT SO_BAO_DANH as SBD, MA_DK as Ma, HO_VA_TEN, NGAY_SINH, 
                    SO_CMT, HANG_GPLX, NOI_DUNG_SH as NOIDUNG, LYTHUYETKT as LT, MOPHONGKT as MP, 
                    HINHKT as H, DUONGKT as Đ, KETQUAKT as KETQUA
                FROM NguoiLX_HoSo 
                ORDER BY SO_BAO_DANH
            """)
            rows = cur.fetchall()
            cols = [d[0] for d in cur.description]
            self.current_columns = cols

            # Xóa dữ liệu cũ
            for old_col in self.data_tree["columns"]:
                try:
                    self.data_tree.heading(old_col, text="")
                except Exception:
                    pass
            self.data_tree.delete(*self.data_tree.get_children())
            
            # Thiết lập cột mới
            self.data_tree["columns"] = cols
            self.data_tree["show"] = "headings"

            # Định nghĩa độ rộng cố định cho từng cột
            fixed_widths = {
                'SBD': 50,
                'Ma': 220,
                'HO_VA_TEN': 180,
                'NGAY_SINH': 90,
                'SO_CMT': 110,
                'HANG_GPLX': 40,
                'NOI_DUNG_SH':30,
                'LT': 50,
                'MP': 50,
                'H': 50,
                'Đ': 50
            }

            # Áp dụng cấu hình cột
            for col in cols:
                width = fixed_widths.get(col, 100)
                anchor = "center" if col in ['SBD', 'LT', 'MP', 'H', 'Đ', 'HANG_GPLX'] else "w"
                
                self.data_tree.heading(col, text=col, anchor="center")
                self.data_tree.column(col, width=width, minwidth=40, anchor=anchor, stretch=False)

            # Thêm dữ liệu vào Treeview
            for r in rows:
                values = []
                for v in r:
                    if isinstance(v, (bytes, bytearray)):
                        values.append("<bytes>")
                    else:
                        values.append("" if v is None else str(v))
                self.data_tree.insert('', 'end', values=values)

            # Cấu hình tag highlight cho các dòng mới thêm
            self.data_tree.tag_configure('highlight', background='#FFFF99')

            self.status_label.config(text=f"Hiển thị {len(rows)} hồ sơ")
            write_log("General", f"Làm mới dữ liệu: Hiển thị {len(rows)} hồ sơ.")

        except Exception as e:
            error_msg = f"LỖI khi tải dữ liệu:\n{e}\n\n{traceback.format_exc()}"
            messagebox.showerror("Lỗi", error_msg)
            write_log("General", error_msg)
    
    
    # Quản lý ĐVHC
    # ---------------------------

    def manage_dvhc(self):
        """Mở cửa sổ quản lý danh mục đơn vị hành chính"""
        dvhc_win = Toplevel(self.root)
        dvhc_win.title("Quản lý Danh mục Đơn vị Hành chính")
        dvhc_win.geometry("1000x600")
        dvhc_win.resizable(True, True)

        # Tiêu đề
        Label(dvhc_win, text="DANH MỤC ĐƠN VỊ HÀNH CHÍNH (INI)",
            font=("Arial", 14, "bold"), fg="blue").pack(pady=10)

        # Frame tìm kiếm
        search_frame = Frame(dvhc_win)
        search_frame.pack(fill=X, padx=10, pady=5)

        Label(search_frame, text="Tìm kiếm:", font=("Arial", 10)).pack(side=LEFT, padx=5)
        search_entry = Entry(search_frame, width=40)
        search_entry.pack(side=LEFT, padx=5)

        # Frame hiển thị
        tree_frame = Frame(dvhc_win)
        tree_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)

        # Scrollbars
        scroll_y = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
        scroll_x = ttk.Scrollbar(tree_frame, orient=tk.HORIZONTAL)

        # Cột theo cấu trúc INI
        columns = ("MA_DVHC", "MA_DV", "MA_DVQL", "TEN_DVHC", "TENNGANGON", "TENDAYDU", "LOAIDVHC")
        
        tree = ttk.Treeview(tree_frame, columns=columns, show="headings",
                        yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)

        scroll_y.config(command=tree.yview)
        scroll_x.config(command=tree.xview)

        # Cấu hình cột
        col_config = {
            "MA_DVHC": {"width": 80, "text": "Mã ĐVHC"},
            "MA_DV": {"width": 60, "text": "Mã ĐV"},
            "MA_DVQL": {"width": 80, "text": "Mã ĐVQL"},
            "TEN_DVHC": {"width": 180, "text": "Tên ĐVHC"},
            "TENNGANGON": {"width": 100, "text": "Tên ngắn"},
            "TENDAYDU": {"width": 320, "text": "Tên đầy đủ"},
            "LOAIDVHC": {"width": 60, "text": "Loại"}
        }

        for col in columns:
            cfg = col_config.get(col, {"width": 100, "text": col})
            tree.heading(col, text=cfg["text"])
            tree.column(col, width=cfg["width"], minwidth=50)

        tree.pack(side=LEFT, fill=BOTH, expand=True)
        scroll_y.pack(side=RIGHT, fill=Y)
        scroll_x.pack(side=BOTTOM, fill=X)

        # Hàm scroll chuột
        def on_mousewheel(event):
            try:
                if tree.winfo_exists():
                    tree.yview_scroll(int(-1*(event.delta/120)), "units")
            except:
                pass
        
        tree.bind("<MouseWheel>", on_mousewheel)

        # Load dữ liệu
        def load_data(filter_text=""):
            tree.delete(*tree.get_children())
            data = ConfigLoader.load_dvhc_data()
            
            filter_lower = filter_text.lower()
            count = 0
            
            for i, item in enumerate(data):
                if filter_text:
                    tendaydu = str(item.get("TENDAYDU", "")).lower()
                    ten_dvhc = str(item.get("TEN_DVHC", "")).lower()
                    tenngangon = str(item.get("TENNGANGON", "")).lower()
                    ma_dvhc = str(item.get("MA_DVHC", ""))
                    
                    if not (filter_lower in tendaydu or filter_lower in ten_dvhc or 
                            filter_lower in tenngangon or filter_lower in ma_dvhc):
                        continue
                
                values = (
                    item.get("MA_DVHC", ""),
                    item.get("MA_DV", ""),
                    item.get("MA_DVQL", ""),
                    item.get("TEN_DVHC", ""),
                    item.get("TENNGANGON", ""),
                    item.get("TENDAYDU", ""),
                    item.get("LOAIDVHC", "")
                )
                tag = "odd" if count % 2 == 0 else "even"
                tree.insert("", "end", values=values, tags=(tag,))
                count += 1
            
            tree.tag_configure("odd", background="#f8f8f8")
            tree.tag_configure("even", background="#ffffff")
            
            status_lbl.config(text=f"Hiển thị {count} / {len(data)} bản ghi | File: {os.path.basename(DM_DVHC_FILE)}")

        def do_search():
            load_data(search_entry.get().strip())

        Button(search_frame, text="🔍 Tìm", command=do_search, width=10).pack(side=LEFT, padx=5)
        Button(search_frame, text="🔄 Tất cả", command=lambda: [search_entry.delete(0, tk.END), load_data()], width=10).pack(side=LEFT, padx=5)

        search_entry.bind("<Return>", lambda e: do_search())

        # Frame nút bấm
        btn_frame = Frame(dvhc_win)
        btn_frame.pack(fill=X, pady=10)

        def reload_file():
            ConfigLoader.reload_dvhc_data()
            load_data()
            messagebox.showinfo("OK", f"Đã reload file:\n{DM_DVHC_FILE}")

        def open_file_location():
            folder = os.path.dirname(DM_DVHC_FILE)
            if os.path.exists(folder):
                os.startfile(folder)
            else:
                messagebox.showwarning("Cảnh báo", f"Thư mục không tồn tại:\n{folder}")

        def show_file_info():
            data = ConfigLoader.load_dvhc_data()
            info = f"📁 File: {DM_DVHC_FILE}\n"
            info += f"📄 Định dạng: INI\n\n"
            info += f"📊 Tổng số bản ghi: {len(data)}\n\n"
            
            if data:
                loai_count = {}
                for item in data:
                    loai = item.get("LOAIDVHC", "Khác") or "Khác"
                    loai_count[loai] = loai_count.get(loai, 0) + 1
                
                info += "📈 Thống kê theo loại:\n"
                loai_names = {
                    "PH": "Phường", "XA": "Xã", "TT": "Thị trấn",
                    "QU": "Quận", "HU": "Huyện", "TP": "Thành phố",
                    "TX": "Thị xã", "TI": "Tỉnh"
                }
                for loai, count in sorted(loai_count.items()):
                    loai_name = loai_names.get(loai, loai)
                    info += f"   • {loai_name} ({loai}): {count}\n"
            
            messagebox.showinfo("Thông tin file ĐVHC", info)

        def on_close_dvhc():
            tree.unbind("<MouseWheel>")
            dvhc_win.destroy()

        Button(btn_frame, text="🔄 Reload", command=reload_file, width=12).pack(side=LEFT, padx=10)
        Button(btn_frame, text="📂 Mở thư mục", command=open_file_location, width=12).pack(side=LEFT, padx=10)
        Button(btn_frame, text="ℹ️ Thông tin", command=show_file_info, width=12).pack(side=LEFT, padx=10)
        Button(btn_frame, text="❌ Đóng", command=on_close_dvhc, width=12, bg="#FFB6C1").pack(side=RIGHT, padx=10)

        dvhc_win.protocol("WM_DELETE_WINDOW", on_close_dvhc)

        # Status
        status_lbl = Label(dvhc_win, text="", relief=tk.SUNKEN, anchor=W)
        status_lbl.pack(fill=X, side=BOTTOM, padx=10, pady=5)

        # Load dữ liệu ban đầu
        load_data()


    # Cập nhật NOI_CT cho dữ liệu hiện có
    # ---------------------------
    
    def update_all_noict_from_dvhc(self):
        """Cập nhật NOI_CT cho tất cả học viên từ DM_DVHC.cfg"""
        # Kiểm tra file ĐVHC
        dvhc_data = ConfigLoader.load_dvhc_data()
        if not dvhc_data:
            messagebox.showerror("Lỗi", f"Không load được file DM_DVHC.cfg\n\nĐường dẫn: {DM_DVHC_FILE}")
            return
        
        # Xác nhận
        if not messagebox.askyesno("Xác nhận", 
            f"Cập nhật cột NOI_CT cho TẤT CẢ học viên?\n\n"
            f"📁 File ĐVHC: {len(dvhc_data)} bản ghi\n\n"
            f"Logic: NOI_CT_MA_DVQL + NOI_CT_MA_DVHC → Tìm MA_DVQL → Lấy TENDAYDU\n\n"
            f"Chỉ những học viên có mã hợp lệ sẽ được cập nhật."):
            return
        
        conn = self.get_connection()
        if not conn:
            return
        
        cur = conn.cursor()
        
        # Lấy tất cả học viên
        cur.execute("SELECT SO_BAO_DANH, HO_VA_TEN, NOI_CT_MA_DVHC, NOI_CT_MA_DVQL, NOI_CT FROM NguoiLX_HoSo")
        rows = cur.fetchall()
        
        if not rows:
            messagebox.showinfo("Thông báo", "Không có dữ liệu học viên để cập nhật.")
            return
        
        updated = 0
        not_found = 0
        no_ma = 0
        log_details = []
        
        for sbd, ho_ten, ma_dvhc, ma_dvql, noi_ct_cu in rows:
            if not ma_dvql or not ma_dvhc:
                no_ma += 1
                continue
            
            # Ghép mã: NOI_CT_MA_DVQL + NOI_CT_MA_DVHC
            ma_dvql_full = str(ma_dvql).strip() + str(ma_dvhc).strip()
            
            # Tìm theo mã DVQL đầy đủ
            tendaydu = ConfigLoader.get_dvhc_by_madvql(ma_dvql_full)
            
            if tendaydu:
                cur.execute("UPDATE NguoiLX_HoSo SET NOI_CT = ? WHERE SO_BAO_DANH = ?", (tendaydu, sbd))
                updated += 1
                log_details.append(f"✓ SBD:{sbd} | {ho_ten} | {ma_dvql_full} → {tendaydu[:40]}...")
            else:
                not_found += 1
                log_details.append(f"✗ SBD:{sbd} | {ho_ten} | Mã:{ma_dvql_full} không tìm thấy")
        
        conn.commit()
        self.refresh_data()
        
        # Ghi log
        write_log("General", f"{'='*60}")
        write_log("General", f"CẬP NHẬT NOI_CT TỪ DM_DVHC.cfg")
        write_log("General", f"Tổng số: {len(rows)} | Thành công: {updated} | Không tìm thấy: {not_found} | Không có mã: {no_ma}")
        write_log("General", f"{'='*60}")
        for detail in log_details[:30]:
            write_log("General", detail)
        if len(log_details) > 30:
            write_log("General", f"... và {len(log_details) - 30} bản ghi khác")
        
        # Thông báo
        msg = f"📍 CẬP NHẬT NOI_CT TỪ DM_DVHC.cfg\n\n"
        msg += f"📊 Tổng số học viên: {len(rows)}\n\n"
        msg += f"✅ Cập nhật thành công: {updated}\n"
        msg += f"❌ Không tìm thấy mã: {not_found}\n"
        msg += f"⚪ Không có mã ĐVHC/ĐVQL: {no_ma}\n"
        
        if not_found > 0:
            msg += f"\n⚠️ Xem chi tiết trong file log."
        
        messagebox.showinfo("Kết quả", msg)


    # ---------------------------
    # view kysh, clear, export minimal
    # ---------------------------
    def view_kysh(self):
        conn = self.get_connection()
        if not conn:
            return
        c = conn.cursor()
        c.execute("SELECT * FROM KY_SH")
        rows = c.fetchall()
        
        if not rows:
            messagebox.showinfo("Kỳ SH", "Chưa có bản ghi KY_SH")
            return
        
        # Lấy tên cột
        c.execute("PRAGMA table_info(KY_SH)")
        cols = [d[1] for d in c.fetchall()]
        
        # Tạo cửa sổ popup
        top = Toplevel(self.root)
        top.title("Thông tin Kỳ Sát Hạch")
        top.geometry("700x500")
        top.resizable(True, True)
        
        # Tiêu đề
        title_label = Label(top, text="THÔNG TIN KỲ SÁT HẠCH", 
                            font=("Arial", 14, "bold"), fg="blue")
        title_label.pack(pady=10)
        
        # Frame chứa bảng
        table_frame = Frame(top)
        table_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)
        
        # Tạo Treeview hiển thị dạng bảng
        tree_scroll_y = ttk.Scrollbar(table_frame, orient=tk.VERTICAL)
        tree_scroll_x = ttk.Scrollbar(table_frame, orient=tk.HORIZONTAL)
        
        tree = ttk.Treeview(table_frame, 
                            columns=("field", "value"),
                            show="headings",
                            yscrollcommand=tree_scroll_y.set,
                            xscrollcommand=tree_scroll_x.set,
                            height=15)
        
        tree_scroll_y.config(command=tree.yview)
        tree_scroll_x.config(command=tree.xview)
        
        # Cấu hình cột
        tree.heading("field", text="Trường dữ liệu")
        tree.heading("value", text="Giá trị")
        tree.column("field", width=200, anchor="w")
        tree.column("value", width=450, anchor="w")
        
        # Mapping tên cột tiếng Việt
        field_names = {
            "MAKYSH": "Mã kỳ sát hạch",
            "MATTSH": "Mã trung tâm sát hạch",
            "NGAYSH": "Ngày sát hạch",
            "GIOSH": "Giờ sát hạch",
            "SOQD": "Số quyết định",
            "NGAYQD": "Ngày quyết định",
            "NGUOIQD": "Người ký quyết định",
            "CHUTICH_HDSH": "Chủ tịch HĐSH",
            "PHOCHUTICH_HDSH": "Phó chủ tịch HĐSH",
            "UV_GD_TTSH": "UV - Giám đốc TTSH",
            "UV_TOTRUONG": "UV - Tổ trưởng",
            "UV_THUKY": "UV - Thư ký",
            "TONGSODK": "Tổng số đăng ký"
        }
        
        # Điền dữ liệu vào bảng (lấy bản ghi đầu tiên)
        row_data = rows[0]
        for i, col in enumerate(cols):
            field_display = field_names.get(col, col)
            value = row_data[i] if i < len(row_data) else ""
            value = "" if value is None else str(value)
            
            # Thêm tag để tô màu xen kẽ
            tag = "odd" if i % 2 == 0 else "even"
            tree.insert("", "end", values=(field_display, value), tags=(tag,))
        
        # Tô màu xen kẽ
        tree.tag_configure("odd", background="#f0f0f0")
        tree.tag_configure("even", background="#ffffff")
        
        # Pack các widget
        tree.pack(side=LEFT, fill=BOTH, expand=True)
        tree_scroll_y.pack(side=RIGHT, fill=Y)
        tree_scroll_x.pack(side=BOTTOM, fill=X)
        
        # Frame nút bấm
        btn_frame = Frame(top)
        btn_frame.pack(fill=X, pady=10)
        
        Button(btn_frame, text="Đóng", command=top.destroy, 
            width=15, height=2, bg="lightgray").pack()
    

    def clear_data(self):
        if messagebox.askyesno("Xác nhận", "Xóa toàn bộ dữ liệu?"):
            conn = self.get_connection()
            c = conn.cursor()
            c.execute("DELETE FROM NguoiLX_HoSo")
            c.execute("DELETE FROM KY_SH")
            conn.commit()
            self.refresh_data()
            write_log("General", "Đã xóa toàn bộ dữ liệu NguoiLX_HoSo và KY_SH.")
            messagebox.showinfo("Xong", "Đã xóa toàn bộ dữ liệu.")

    #Hàm tìm kiếm SBD
    
    def search_sbd(self):
        sbd = self.entry_search_sbd.get().strip()

        if not sbd:
            messagebox.showwarning("Thiếu dữ liệu", "Vui lòng nhập SBD.")
            return

        conn = self.get_connection()
        cur = conn.cursor()

        # Sử dụng cùng câu truy vấn như refresh_data() để đảm bảo hiển thị đúng cột
        cur.execute("""
            SELECT SO_BAO_DANH as SBD, MA_DK as Ma, HO_VA_TEN as HoVaTen, NGAY_SINH as NgaySinh, 
                SO_CMT as CMT, HANG_GPLX as Hang, NOI_DUNG_SH as NoiDung, LYTHUYETKT as LT, MOPHONGKT as MP, 
                HINHKT as H, DUONGKT as Đ, KETQUAKT as KetQua
            FROM NguoiLX_HoSo 
            WHERE SO_BAO_DANH = ?
        """, (sbd,))
        rows = cur.fetchall()

        # Xóa tree + hiển thị dòng tìm thấy
        self.data_tree.delete(*self.data_tree.get_children())

        if not rows:
            messagebox.showinfo("Kết quả", f"Không tìm thấy SBD {sbd}")
            return

        for r in rows:
            self.data_tree.insert("", "end", values=r)

        self.status_label.config(text=f"Tìm thấy {len(rows)} kết quả cho SBD {sbd}")

    #Hàm tìm kiếm Họ và tên

    def search_ho_ten(self):
        """Tìm kiếm học viên theo họ tên (không phân biệt hoa/thường, hỗ trợ tiếng Việt)"""
        ten = self.entry_search_name.get().strip()

        if not ten:
            messagebox.showwarning("Thiếu dữ liệu", "Vui lòng nhập họ tên cần tìm.")
            return

        conn = self.get_connection()
        if not conn:
            return
            
        cur = conn.cursor()

        # Lấy tất cả dữ liệu
        cur.execute("""
            SELECT SO_BAO_DANH as SBD, MA_DK as Ma, HO_VA_TEN as HoVaTen, NGAY_SINH as NgaySinh, 
                SO_CMT as CMT, HANG_GPLX as Hang, NOI_DUNG_SH as NoiDung, LYTHUYETKT as LT, MOPHONGKT as MP, 
                HINHKT as H, DUONGKT as Đ, KETQUAKT as KetQua
            FROM NguoiLX_HoSo 
            ORDER BY SO_BAO_DANH
        """)
        
        all_rows = cur.fetchall()
        
        # Chuyển từ khóa tìm kiếm thành chữ thường
        ten_lower = ten.lower()
        
        # Lọc kết quả bằng Python (hỗ trợ Unicode tiếng Việt)
        matched_rows = []
        for row in all_rows:
            ho_va_ten = row[2]  # Cột HO_VA_TEN ở index 2
            if ho_va_ten and ten_lower in ho_va_ten.lower():
                matched_rows.append(row)

        # Xóa dữ liệu cũ trong Treeview
        self.data_tree.delete(*self.data_tree.get_children())

        if not matched_rows:
            messagebox.showinfo("Kết quả", f"Không tìm thấy học viên có tên chứa '{ten}'")
            self.status_label.config(text=f"Không tìm thấy kết quả cho '{ten}'")
            return

        # Hiển thị kết quả
        for r in matched_rows:
            values = []
            for v in r:
                if isinstance(v, (bytes, bytearray)):
                    values.append("<bytes>")
                else:
                    values.append("" if v is None else str(v))
            self.data_tree.insert("", "end", values=values)

        self.status_label.config(text=f"Tìm thấy {len(matched_rows)} kết quả cho '{ten}'")
        write_log("General", f"Tìm kiếm họ tên '{ten}': {len(matched_rows)} kết quả")


    # hàm kiểm tra nội dung
    
    def check_noidung_thi(self, sbd, db_column):
        """
        Kiểm tra học viên có phải thi môn đó không dựa trên NOI_DUNG_SH
        
        Args:
            sbd: Số báo danh
            db_column: Tên cột trong DB (LYTHUYETKT, MOPHONGKT, HINHKT, DUONGKT)
        
        Returns:
            tuple: (phai_thi: bool, ho_ten: str, noi_dung_sh: str, message: str)
        """
        conn = self.get_connection()
        if not conn:
            return True, "", "", "Không kết nối được CSDL - Cho phép nhập"
        
        cur = conn.cursor()
        cur.execute("SELECT HO_VA_TEN, NOI_DUNG_SH, HANG_GPLX FROM NguoiLX_HoSo WHERE SO_BAO_DANH = ?", (sbd,))
        row = cur.fetchone()
        
        if not row:
            return False, "", "", f"Không tìm thấy SBD: {sbd}"
        
        ho_ten = row[0] or ""
        noi_dung_sh = str(row[1] or "").strip()
        hang_gplx = row[2] or ""
        
        # ========== MAPPING CỘT DB → KEY TRONG NOIDUNG_MAP ==========
        column_to_key = {
            "LYTHUYETKT": "LT",
            "MOPHONGKT": "MP",
            "HINHKT": "H",
            "DUONGKT": "D"
        }
        
        column_names = {
            "LYTHUYETKT": "Lý thuyết",
            "MOPHONGKT": "Mô phỏng",
            "HINHKT": "Hình",
            "DUONGKT": "Đường"
        }
        
        # Lấy key tương ứng
        mon_key = column_to_key.get(db_column)
        ten_mon = column_names.get(db_column, db_column)
        
        if not mon_key:
            return True, ho_ten, noi_dung_sh, f"Không xác định môn {db_column}"
        
        # ========== KIỂM TRA NỘI DUNG ==========
        
        # Nếu NOI_DUNG_SH trống → Cho phép tất cả
        if not noi_dung_sh:
            return True, ho_ten, "", f"Không có ND thi - Cho phép nhập {ten_mon}"
        
        # Nếu NOI_DUNG_SH không có trong bảng → Cho phép (an toàn)
        if noi_dung_sh not in NOIDUNG_MAP:
            return True, ho_ten, noi_dung_sh, f"ND '{noi_dung_sh}' không xác định - Cho phép nhập"
        
        # Lấy thông tin các môn cần thi
        required = NOIDUNG_MAP[noi_dung_sh]
        phai_thi = required.get(mon_key, False)
        
        # Lấy mô tả nội dung
        mota = NOIDUNG_MOTA.get(noi_dung_sh, "")
        
        # Tạo danh sách các môn phải thi
        ds_mon_phai_thi = []
        if required.get("LT"): ds_mon_phai_thi.append("LT")
        if required.get("MP"): ds_mon_phai_thi.append("MP")
        if required.get("H"): ds_mon_phai_thi.append("H")
        if required.get("D"): ds_mon_phai_thi.append("D")
        
        ds_text = "+".join(ds_mon_phai_thi) if ds_mon_phai_thi else "Không có"
        
        # Tạo message
        if phai_thi:
            message = f"✅ HV phải thi {ten_mon}\nND: {noi_dung_sh} ({mota})"
        else:
            message = f"❌ HV KHÔNG phải thi {ten_mon}\n\nND: {noi_dung_sh} ({mota})\nChỉ thi: {ds_text}"
        
        return phai_thi, ho_ten, noi_dung_sh, message


    # ---------- CHỨC NĂNG SỬA DỮ LIỆU ----------
    def edit_item(self, event):
        """Mở cửa sổ chỉnh sửa với TẤT CẢ các cột trong bảng NguoiLX_HoSo"""
        # Lấy item được chọn
        selected = self.data_tree.focus()
        if not selected:
            return
        
        # Lấy giá trị từ dòng được chọn (chỉ có 10 cột hiển thị)
        values = self.data_tree.item(selected, "values")
        if not values:
            messagebox.showwarning("Lỗi", "Không có dữ liệu để chỉnh sửa.")
            return

        # Lấy SBD từ cột đầu tiên
        sbd_value = values[0] if len(values) > 0 else None
        if not sbd_value:
            messagebox.showwarning("Lỗi", "Không tìm thấy SBD.")
            return

        # Query lấy TẤT CẢ thông tin từ database
        conn = self.get_connection()
        if not conn:
            return
        
        cur = conn.cursor()
        
        # Lấy tên tất cả các cột trong bảng
        cur.execute("PRAGMA table_info(NguoiLX_HoSo)")
        all_columns = [col[1] for col in cur.fetchall()]
        
        # Lấy dữ liệu của học viên theo SBD
        cur.execute("SELECT * FROM NguoiLX_HoSo WHERE SO_BAO_DANH = ?", (sbd_value,))
        row_data = cur.fetchone()
        
        if not row_data:
            messagebox.showwarning("Lỗi", f"Không tìm thấy học viên có SBD: {sbd_value}")
            return

        # Lấy họ tên để hiển thị tiêu đề
        ho_ten_idx = all_columns.index('HO_VA_TEN') if 'HO_VA_TEN' in all_columns else -1
        ho_ten = row_data[ho_ten_idx] if ho_ten_idx >= 0 else 'Học viên'

        # Tạo cửa sổ popup
        popup_width = 1000
        popup_height = 650
        edit_win = Toplevel(self.root)
        edit_win.title(f"Chỉnh sửa: {ho_ten} (SBD: {sbd_value})")
        edit_win.geometry(f"{popup_width}x{popup_height}+50+30")
        edit_win.resizable(True, True)

        # Tiêu đề
        title_frame = Frame(edit_win)
        title_frame.pack(fill="x", padx=10, pady=5)
        Label(title_frame, text=f"THÔNG TIN HỌC VIÊN: {ho_ten}", 
            font=("Arial", 14, "bold"), fg="blue").pack()
        Label(title_frame, text=f"Số báo danh: {sbd_value} | Tổng số trường: {len(all_columns)}", 
            font=("Arial", 10)).pack()

        # Khung chính với scrollbar
        frame_main = Frame(edit_win)
        frame_main.pack(fill=BOTH, expand=True, padx=10, pady=5)

        # Canvas và Scrollbar
        canvas = Canvas(frame_main)
        scrollbar_y = Scrollbar(frame_main, orient="vertical", command=canvas.yview)
        scrollbar_x = Scrollbar(frame_main, orient="horizontal", command=canvas.xview)
        scroll_frame = Frame(canvas)

        scroll_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scroll_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar_y.set, xscrollcommand=scrollbar_x.set)
        
        # Pack scrollbars
        scrollbar_y.pack(side=RIGHT, fill=Y)
        scrollbar_x.pack(side=BOTTOM, fill=X)
        canvas.pack(side=LEFT, fill=BOTH, expand=True)

        # Hỗ trợ scroll bằng chuột
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", on_mousewheel)
        
        # Đóng sự kiện khi cửa sổ đóng
        def on_close():
            canvas.unbind_all("<MouseWheel>")
            edit_win.destroy()
        edit_win.protocol("WM_DELETE_WINDOW", on_close)

        # Mapping tên cột tiếng Việt
        column_labels = {
            "SO_TT": "Số TT",
            "MA_DK": "Mã đăng ký",
            "HO_TEN_DEM": "Họ tên đệm",
            "TEN": "Tên",
            "HO_VA_TEN": "Họ và tên",
            "GIOI_TINH": "Giới tính",
            "NGAY_SINH": "Ngày sinh",
            "MA_QUOC_TICH": "Mã quốc tịch",
            "NOI_CT": "Nơi cư trú",
            "NOI_CT_MA_DVHC": "Mã ĐVHC nơi CT",
            "NOI_CT_MA_DVQL": "Mã ĐVQL nơi CT",
            "SO_CMT": "Số CMT/CCCD",
            "SO_HO_SO": "Số hồ sơ",
            "MA_KY_SH": "Mã kỳ sát hạch",
            "SO_BAO_DANH": "Số báo danh",
            "MA_CSDT": "Mã CSĐT",
            "MA_TTSH": "Mã TTSH",
            "MA_SO_GTVT": "Mã Sở GTVT",
            "GIAY_CNSK": "Giấy CN sức khỏe",
            "HANG_GPLX": "Hạng GPLX",
            "SO_GPLX_DA_CO": "Số GPLX đã có",
            "HANG_GPLX_DA_CO": "Hạng GPLX đã có",
            "DVQL_GPLX_DACO": "ĐVQL GPLX đã có",
            "NGAY_HH_GPLX_DACO": "Ngày HH GPLX đã có",
            "SO_NAM_LAIXE": "Số năm lái xe",
            "SO_KM_ANTOAN": "Số km an toàn",
            "SO_GIAY_CNTN": "Số giấy CNTN",
            "SO_CCN": "Số CCN",
            "NOI_DUNG_SH": "Nội dung SH",
            "LY_DO_SH": "Lý do SH",
            "KET_QUA_SH": "Kết quả SH",
            "KQ_SH_LYTHUYET": "KQ Lý thuyết",
            "KQ_SH_MOPHONG": "KQ Mô phỏng",
            "KQ_SH_HINH": "KQ Hình",
            "KQ_SH_DUONG": "KQ Đường",
            "GHI_CHU_SH": "Ghi chú SH",
            "ANH_CHAN_DUNG": "Ảnh chân dung",
            "NGAY_TT_GPLX_DACO": "Ngày TT GPLX đã có",
            "MA_KHOA_HOC": "Mã khóa học",
            "SO_QD_SH": "Số QĐ SH",
            "NGAY_QD_SH": "Ngày QĐ SH",
            "NGUOI_QD_SH": "Người QĐ SH",
            "CHAT_LUONG_ANH": "Chất lượng ảnh",
            "LYTHUYETKT": "Lý thuyết KT",
            "MOPHONGKT": "Mô phỏng KT",
            "HINHKT": "Hình KT",
            "DUONGKT": "Đường KT"
        }

        entries = {}
        num_cols = 2  # Số cột hiển thị (2 cặp label-entry)
        mid = (len(all_columns) + num_cols - 1) // num_cols

        for i, col in enumerate(all_columns):
            current_value = row_data[i] if i < len(row_data) else ""
            # Chuyển None thành chuỗi rỗng
            if current_value is None:
                current_value = ""
            else:
                current_value = str(current_value)
            
            row_idx = i % mid
            col_start = (i // mid) * 2
            
            # Label với tên tiếng Việt
            label_text = column_labels.get(col, col)
            lbl = Label(scroll_frame, text=f"{label_text}:", width=22, anchor=W, font=("Arial", 9))
            lbl.grid(row=row_idx, column=col_start, padx=5, pady=3, sticky=W)
            
            # Entry
            entry_width = 35
            e = Entry(scroll_frame, width=entry_width, font=("Arial", 9))
            e.insert(0, current_value)
            e.grid(row=row_idx, column=col_start+1, padx=5, pady=3, sticky=W+E)
            entries[col] = e
            
            # Khóa cột SO_BAO_DANH
            if col == 'SO_BAO_DANH':
                e.config(state=DISABLED, disabledbackground="#f0f0f0")
            
            # Tô màu các trường quan trọng
            if col in ['HO_VA_TEN', 'NGAY_SINH', 'SO_CMT', 'HANG_GPLX']:
                lbl.config(fg="blue", font=("Arial", 9, "bold"))
            
            # Tô màu các trường kết quả
            if col in ['KQ_SH_LYTHUYET', 'KQ_SH_MOPHONG', 'KQ_SH_HINH', 'KQ_SH_DUONG', 
                    'LYTHUYETKT', 'MOPHONGKT', 'HINHKT', 'DUONGKT']:
                lbl.config(fg="green", font=("Arial", 9, "bold"))

        # Frame nút bấm
        btn_frame = Frame(edit_win)
        btn_frame.pack(fill=X, side=BOTTOM, pady=10)
        
        Button(btn_frame, text="💾 Lưu thay đổi",
            command=lambda: self.save_changes_full(entries, all_columns, sbd_value, edit_win),
            bg="#90EE90", width=18, height=2, font=("Arial", 10, "bold")).pack(side=LEFT, padx=20)
        
        Button(btn_frame, text="❌ Hủy bỏ",
            command=on_close,
            bg="#FFB6C1", width=18, height=2, font=("Arial", 10, "bold")).pack(side=RIGHT, padx=20)
        
        # Cấu hình cột giãn
        scroll_frame.columnconfigure(1, weight=1)
        scroll_frame.columnconfigure(3, weight=1)



    # Thêm các hàm xử lý


    def focus_next_input(self):
        """Focus vào ô tiếp theo hoặc nút đầu tiên"""
        # Có thể thêm logic focus vào nút LT Đạt nếu muốn
        pass
    
    
    # update_result_by_sbd
    def update_result_by_sbd(self, db_column, result_value):
        """
        Cập nhật kết quả thi cho học viên theo SBD đã nhập
        Có kiểm tra NOI_DUNG_SH trước khi cập nhật
        """
        # Lấy SBD từ ô nhập
        sbd = self.entry_input_sbd.get().strip()
        
        if not sbd:
            messagebox.showwarning("Thiếu dữ liệu", "Vui lòng nhập SBD trước khi chọn kết quả!")
            self.entry_input_sbd.focus_set()
            return
        
        conn = self.get_connection()
        if not conn:
            return
        
        cur = conn.cursor()
        
        # Kiểm tra SBD có tồn tại không
        cur.execute("SELECT HO_VA_TEN, NOI_DUNG_SH, HANG_GPLX FROM NguoiLX_HoSo WHERE SO_BAO_DANH = ?", (sbd,))
        row = cur.fetchone()
        
        if not row:
            messagebox.showerror("Lỗi", f"Không tìm thấy học viên có SBD: {sbd}")
            self.entry_input_sbd.focus_set()
            self.entry_input_sbd.select_range(0, tk.END)
            return
        
        ho_ten = row[0]
        noi_dung_sh = str(row[1] or "").strip()
        hang_gplx = row[2] or ""
        
        # Mapping tên môn
        column_names = {
            "LYTHUYETKT": "Lý thuyết",
            "MOPHONGKT": "Mô phỏng",
            "HINHKT": "Hình",
            "DUONGKT": "Đường"
        }
        
        result_names = {
            "DA": "Đạt ✅",
            "RO": "Trượt ❌",
            "VA": "Vắng ⬜"
        }
        
        mon_thi = column_names.get(db_column, db_column)
        ket_qua = result_names.get(result_value, result_value)
        
        # ========== KIỂM TRA NỘI DUNG THI ==========
        phai_thi, _, nd_sh, message = self.check_noidung_thi(sbd, db_column)
        
        # ========== NẾU KHÔNG PHẢI THI ==========
        if not phai_thi:
            # Lấy mô tả nội dung
            mota = NOIDUNG_MOTA.get(nd_sh, "")
            
            # Lấy danh sách môn phải thi
            ds_mon = []
            if nd_sh in NOIDUNG_MAP:
                req = NOIDUNG_MAP[nd_sh]
                if req.get("LT"): ds_mon.append("Lý thuyết")
                if req.get("MP"): ds_mon.append("Mô phỏng")
                if req.get("H"): ds_mon.append("Hình")
                if req.get("D"): ds_mon.append("Đường")
            ds_text = ", ".join(ds_mon) if ds_mon else "Không xác định"
            
            # Hiển thị cảnh báo
            confirm = messagebox.askyesno(
                "⚠️ Cảnh báo - Không phải thi môn này",
                f"👤 Học viên: {ho_ten}\n"
                f"📋 SBD: {sbd}\n"
                f"🎫 Hạng: {hang_gplx}\n"
                f"📝 Nội dung SH: {nd_sh} ({mota})\n\n"
                f"❌ HV KHÔNG phải thi môn: {mon_thi}\n"
                f"✅ Chỉ thi: {ds_text}\n\n"
                f"Bạn có muốn vẫn nhập kết quả không?",
                icon='warning'
            )
            
            if not confirm:
                self.status_label.config(
                    text=f"⚠️ SBD {sbd}: Không phải thi {mon_thi} (ND: {nd_sh})"
                )
                self.entry_input_sbd.focus_set()
                return
        
        # ========== CẬP NHẬT KẾT QUẢ ==========
        try:
            cur.execute(f"UPDATE NguoiLX_HoSo SET {db_column} = ? WHERE SO_BAO_DANH = ?", 
                    (result_value, sbd))
            conn.commit()
            
            # Hiển thị thông báo
            self.status_label.config(
                text=f"✅ SBD {sbd} - {ho_ten}: {mon_thi} = {ket_qua}"
            )
            
            # Ghi log
            write_log("General", f"Cập nhật KQ: SBD={sbd}, {db_column}={result_value}, ND={nd_sh}")
            
            # Hiển thị dòng trong Treeview
            self.show_sbd_in_treeview(sbd)
            
            # Focus lại ô nhập
            self.entry_input_sbd.focus_set()
            
        except Exception as e:
            error_msg = f"Lỗi cập nhật kết quả: {e}"
            messagebox.showerror("Lỗi", error_msg)
            write_log("General", f"LỖI: {error_msg}")
    



    #hàm hiển thị thông tin nội dung thi khi nhập SBD
    
    def on_sbd_input_change(self, event=None):
        """Hiển thị tên học viên và nội dung thi khi nhập SBD"""
        sbd = self.entry_input_sbd.get().strip()
        
        if not sbd:
            self.lbl_student_name.config(text="")
            return
        
        conn = self.get_connection()
        if not conn:
            return
        
        cur = conn.cursor()
        cur.execute("""
            SELECT HO_VA_TEN, HANG_GPLX, NOI_DUNG_SH 
            FROM NguoiLX_HoSo 
            WHERE SO_BAO_DANH = ?
        """, (sbd,))
        row = cur.fetchone()
        
        if row:
            ho_ten = row[0] or ""
            hang = row[1] or ""
            noi_dung = str(row[2] or "").strip()
            
            # Lấy mô tả và danh sách môn thi
            if noi_dung and noi_dung in NOIDUNG_MAP:
                mota = NOIDUNG_MOTA.get(noi_dung, "")
                req = NOIDUNG_MAP[noi_dung]
                ds_mon = []
                if req.get("LT"): ds_mon.append("LT")
                if req.get("MP"): ds_mon.append("MP")
                if req.get("H"): ds_mon.append("H")
                if req.get("D"): ds_mon.append("D")
                ds_text = "+".join(ds_mon)
                
                self.lbl_student_name.config(
                    text=f"👤 {ho_ten} | Hạng {hang}\n📝 ND: {noi_dung} → Thi: {ds_text}", 
                    foreground="blue"
                )
            elif noi_dung:
                # Có nội dung nhưng không trong bảng
                self.lbl_student_name.config(
                    text=f"👤 {ho_ten} | Hạng {hang}\n⚠️ ND: {noi_dung} (không xác định)", 
                    foreground="orange"
                )
            else:
                # Không có nội dung
                self.lbl_student_name.config(
                    text=f"👤 {ho_ten} | Hạng {hang}\n📝 ND: (chưa có)", 
                    foreground="blue"
                )
        else:
            self.lbl_student_name.config(text="⚠️ Không tìm thấy SBD", foreground="red")



    def show_sbd_in_treeview(self, sbd):
        """Hiển thị dòng của SBD trong Treeview và highlight"""
        conn = self.get_connection()
        if not conn:
            return
        
        cur = conn.cursor()
        
        # Lấy dữ liệu của SBD
        cur.execute("""
            SELECT SO_BAO_DANH as SBD, MA_DK as Ma, HO_VA_TEN as HoVaTen, NGAY_SINH as NgaySinh, 
                SO_CMT as CMT, HANG_GPLX as Hang, NOI_DUNG_SH as NoiDung, LYTHUYETKT as LT, MOPHONGKT as MP, 
                HINHKT as H, DUONGKT as Đ, KETQUAKT as KetQua
            FROM NguoiLX_HoSo 
            WHERE SO_BAO_DANH = ?
        """, (sbd,))
        row = cur.fetchone()
        
        if not row:
            return
        
        # Tìm và xóa dòng cũ của SBD này trong Treeview (nếu có)
        for item in self.data_tree.get_children():
            values = self.data_tree.item(item, "values")
            if values and str(values[0]) == str(sbd):
                self.data_tree.delete(item)
                break
        
        # Thêm dòng mới vào đầu Treeview
        values = []
        for v in row:
            if isinstance(v, (bytes, bytearray)):
                values.append("<bytes>")
            else:
                values.append("" if v is None else str(v))
        
        # Insert vào đầu và highlight
        new_item = self.data_tree.insert('', 0, values=values, tags=('highlight',))
        
        # Cấu hình tag highlight (màu vàng nhạt)
        self.data_tree.tag_configure('highlight', background='#FFFF99')
        
        # Scroll lên đầu và chọn dòng vừa thêm
        self.data_tree.see(new_item)
        self.data_tree.selection_set(new_item)
        self.data_tree.focus(new_item)



    # hàm đối sánh kết quả thi
    def compare_results(self):
        """
        Đối sánh kết quả thi:
        1. Đối chiếu NOI_DUNG_SH với KQ_SH_* - điền VA nếu thiếu, báo lỗi nếu thừa
        2. Đối chiếu KQ_SH_* với *KT - báo lỗi nếu không trùng
        """
        conn = self.get_connection()
        if not conn:
            return
        
        # Sử dụng NOIDUNG_MAP đã định nghĩa toàn cục
        
        try:
            cur = conn.cursor()
            
            # ... phần còn lại giữ nguyên, thay NOIDUNG_MAP cục bộ bằng biến toàn cục ...
            cur = conn.cursor()
            
            # Lấy tất cả học viên
            cur.execute("""
                SELECT SO_BAO_DANH as SBD, MA_DK as Ma, HO_VA_TEN as HoVaTen, NGAY_SINH as NgaySinh, 
                SO_CMT as CMT, HANG_GPLX as Hang, NOI_DUNG_SH as NoiDung, LYTHUYETKT as LT, MOPHONGKT as MP, 
                HINHKT as H, DUONGKT as Đ, KETQUAKT as KetQua
                FROM NguoiLX_HoSo
                ORDER BY SO_BAO_DANH
            """)
            rows = cur.fetchall()
            
            if not rows:
                messagebox.showinfo("Thông báo", "Không có dữ liệu học viên.")
                return
            
            # Thống kê
            total = len(rows)
            updated_va = 0
            error_noidung = 0
            error_kt = 0
            no_noidung = 0
            ok_count = 0
            
            # Danh sách chi tiết
            list_updated_va = []      # Đã điền VA
            list_error_noidung = []   # Lỗi thừa kết quả (không khớp nội dung)
            list_error_kt = []        # Lỗi không khớp với KT
            list_no_noidung = []      # Không có mã nội dung
            
            for row in rows:
                sbd, ho_ten, noi_dung_sh, kq_lt, kq_mp, kq_h, kq_d, lt_kt, mp_kt, h_kt, d_kt = row
                
                # Chuẩn hóa giá trị
                noi_dung_sh = str(noi_dung_sh or "").strip()
                kq_lt = str(kq_lt or "").strip().upper()
                kq_mp = str(kq_mp or "").strip().upper()
                kq_h = str(kq_h or "").strip().upper()
                kq_d = str(kq_d or "").strip().upper()
                lt_kt = str(lt_kt or "").strip().upper()
                mp_kt = str(mp_kt or "").strip().upper()
                h_kt = str(h_kt or "").strip().upper()
                d_kt = str(d_kt or "").strip().upper()
                
                has_error = False
                error_details = []
                va_updates = []
                
                # ========== 1. ĐỐI CHIẾU NOI_DUNG_SH VỚI KQ_SH_* ==========
                if not noi_dung_sh or noi_dung_sh not in NOIDUNG_MAP:
                    no_noidung += 1
                    list_no_noidung.append({
                        "sbd": sbd,
                        "ho_ten": ho_ten,
                        "noi_dung": noi_dung_sh or "(trống)",
                        "kq_lt": kq_lt, "kq_mp": kq_mp, "kq_h": kq_h, "kq_d": kq_d
                    })
                    continue
                
                required = NOIDUNG_MAP[noi_dung_sh]
                
                # Kiểm tra từng môn
                # Lý thuyết
                if required["LT"]:
                    if not kq_lt:
                        va_updates.append("KQ_SH_LYTHUYET")
                        kq_lt = "VA"
                else:
                    if kq_lt and kq_lt in ["DA", "RO"]:
                        has_error = True
                        error_details.append(f"LT thừa ({kq_lt})")
                
                # Mô phỏng
                if required["MP"]:
                    if not kq_mp:
                        va_updates.append("KQ_SH_MOPHONG")
                        kq_mp = "VA"
                else:
                    if kq_mp and kq_mp in ["DA", "RO"]:
                        has_error = True
                        error_details.append(f"MP thừa ({kq_mp})")
                
                # Hình
                if required["H"]:
                    if not kq_h:
                        va_updates.append("KQ_SH_HINH")
                        kq_h = "VA"
                else:
                    if kq_h and kq_h in ["DA", "RO"]:
                        has_error = True
                        error_details.append(f"H thừa ({kq_h})")
                
                # Đường
                if required["D"]:
                    if not kq_d:
                        va_updates.append("KQ_SH_DUONG")
                        kq_d = "VA"
                else:
                    if kq_d and kq_d in ["DA", "RO"]:
                        has_error = True
                        error_details.append(f"D thừa ({kq_d})")
                
                # Ghi nhận lỗi nội dung
                if has_error:
                    error_noidung += 1
                    list_error_noidung.append({
                        "sbd": sbd,
                        "ho_ten": ho_ten,
                        "noi_dung": noi_dung_sh,
                        "kq_lt": kq_lt, "kq_mp": kq_mp, "kq_h": kq_h, "kq_d": kq_d,
                        "error": ", ".join(error_details)
                    })
                
                # Cập nhật VA vào DB
                if va_updates:
                    update_sql = ", ".join([f"{col} = 'VA'" for col in va_updates])
                    cur.execute(f"UPDATE NguoiLX_HoSo SET {update_sql} WHERE SO_BAO_DANH = ?", (sbd,))
                    updated_va += 1
                    list_updated_va.append({
                        "sbd": sbd,
                        "ho_ten": ho_ten,
                        "noi_dung": noi_dung_sh,
                        "updated_cols": ", ".join(va_updates)
                    })
                
                # ========== 2. ĐỐI CHIẾU KQ_SH_* VỚI *KT ==========
                kt_errors = []
                
                # So sánh LT
                if kq_lt and lt_kt and kq_lt != lt_kt:
                    kt_errors.append(f"LT: {kq_lt} ≠ {lt_kt}")
                
                # So sánh MP
                if kq_mp and mp_kt and kq_mp != mp_kt:
                    kt_errors.append(f"MP: {kq_mp} ≠ {mp_kt}")
                
                # So sánh H
                if kq_h and h_kt and kq_h != h_kt:
                    kt_errors.append(f"H: {kq_h} ≠ {h_kt}")
                
                # So sánh D
                if kq_d and d_kt and kq_d != d_kt:
                    kt_errors.append(f"D: {kq_d} ≠ {d_kt}")
                
                if kt_errors:
                    error_kt += 1
                    list_error_kt.append({
                        "sbd": sbd,
                        "ho_ten": ho_ten,
                        "kq_lt": kq_lt, "kq_mp": kq_mp, "kq_h": kq_h, "kq_d": kq_d,
                        "lt_kt": lt_kt, "mp_kt": mp_kt, "h_kt": h_kt, "d_kt": d_kt,
                        "error": ", ".join(kt_errors)
                    })
                
                # Đếm OK
                if not has_error and not kt_errors and not va_updates:
                    ok_count += 1
            
            conn.commit()
            self.refresh_data()
            
            # ========== HIỂN THỊ KẾT QUẢ ==========
            self.show_compare_result_popup(
                total=total,
                ok_count=ok_count,
                updated_va=updated_va,
                error_noidung=error_noidung,
                error_kt=error_kt,
                no_noidung=no_noidung,
                list_updated_va=list_updated_va,
                list_error_noidung=list_error_noidung,
                list_error_kt=list_error_kt,
                list_no_noidung=list_no_noidung,
                noidung_map=NOIDUNG_MAP
            )
            
            # Ghi log
            write_log("General", f"ĐỐI SÁNH KẾT QUẢ: Tổng {total} | OK {ok_count} | VA {updated_va} | Lỗi ND {error_noidung} | Lỗi KT {error_kt}")
            
        except Exception as e:
            error_msg = f"Lỗi đối sánh: {e}\n{traceback.format_exc()}"
            messagebox.showerror("Lỗi", error_msg)
            write_log("General", error_msg)




    def save_changes_full(self, entries, all_columns, sbd_value, edit_win):
        """Lưu tất cả thay đổi vào database"""
        conn = self.get_connection()
        if not conn:
            return
        
        try:
            c = conn.cursor()
            
            # Chuẩn bị câu lệnh SQL UPDATE
            set_clauses = []
            data = []
            
            for col in all_columns:
                if col == 'SO_BAO_DANH':
                    # Bỏ qua cột SBD (không update)
                    continue
                
                entry = entries[col]
                value = entry.get()
                
                # Chuyển chuỗi rỗng thành None cho các trường số
                if value == "":
                    value = None
                
                set_clauses.append(f'{col}=?')
                data.append(value)
            
            # Thêm SBD cho điều kiện WHERE
            data.append(sbd_value)
            
            sql = f"UPDATE NguoiLX_HoSo SET {', '.join(set_clauses)} WHERE SO_BAO_DANH=?"
            
            c.execute(sql, data)
            conn.commit()
            
            write_log("General", f"Đã lưu chỉnh sửa đầy đủ cho SBD: {sbd_value}.")
            messagebox.showinfo("Thành công", f"Đã lưu chỉnh sửa cho SBD: {sbd_value}")
            edit_win.destroy()
            self.refresh_data()
            
        except Exception as e:
            error_msg = f"LỖI khi lưu chỉnh sửa: {e}\n\n{traceback.format_exc()}"
            messagebox.showerror("Lỗi DB", error_msg)
            write_log("General", error_msg)

    def save_changes(self, entries, sbd_value, edit_win):
        conn = self.get_connection()
        if not conn:
            return
        
        try:
            c = conn.cursor()
            cols = self.current_columns
            
            # Chuẩn bị câu lệnh SQL
            set_clauses = [f'{col}=?' for col in cols]
            sql = f"UPDATE NguoiLX_HoSo SET {', '.join(set_clauses)} WHERE SO_BAO_DANH=?"
            
            # Chuẩn bị dữ liệu: Lấy giá trị từ Entry, kích hoạt lại Entry SBD nếu bị Disable
            data = []
            for col in cols:
                entry = entries[col]
                if col == 'SO_BAO_DANH' and entry.cget('state') == DISABLED:
                    # Vì SBD bị disabled, cần lấy giá trị gốc đã truyền vào hàm (đã được xác thực)
                    data.append(sbd_value)
                else:
                    data.append(entry.get())
            
            # Thêm SBD cho điều kiện WHERE
            data.append(sbd_value) 
            
            c.execute(sql, data)
            conn.commit()
            write_log("General", f"Đã lưu chỉnh sửa cho SBD: {sbd_value}.")
            messagebox.showinfo("OK", f"Đã lưu chỉnh sửa cho SBD: {sbd_value}.")
            edit_win.destroy()
            self.refresh_data()
            
        except Exception as e:
            error_msg = f"LỖI khi lưu chỉnh sửa: {e}\n\n{traceback.format_exc()}"
            messagebox.showerror("Lỗi DB", error_msg)
            write_log("General", error_msg)


    # ---------- XUẤT FILE WORD ----------
    def export_to_word(self):
        """Xuất thẻ dự thi ra Word - 5 hàng x 2 cột (10 thẻ/trang) - KHÔNG TRANG TRẮNG"""
        try:
            conn = sqlite3.connect(self.db_path)
            c = conn.cursor()
            
            # Lấy ngày thi
            c.execute("SELECT NGAYSH FROM KY_SH")
            row = c.fetchone()
            ngay_thi = row[0] if row else ""
            
            # Lấy danh sách học viên
            c.execute("SELECT SO_BAO_DANH, HO_VA_TEN, ANH_CHAN_DUNG, HANG_GPLX FROM NguoiLX_HoSo ORDER BY SO_BAO_DANH")
            hoc_vien_list = c.fetchall()
            conn.close()
            
            if not hoc_vien_list:
                messagebox.showwarning("Cảnh báo", "Không có dữ liệu để xuất!")
                return
            
            # Chọn file lưu
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx")],
                initialfile="The_Du_Thi.docx"
            )
            if not save_path:
                return
            
            doc = Document()
            
            # ========== KÍCH THƯỚC ==========
            CARD_WIDTH = 3.62
            CARD_HEIGHT = 2.2
            NUM_ROWS = 5
            NUM_COLS = 2
            CARDS_PER_PAGE = NUM_ROWS * NUM_COLS
            
            COL_IMG = 1.2
            COL_INFO = 2.42
            ROW_HEADER = 0.7
            ROW_BODY = 1.5
            IMG_WIDTH = 2 / 2.54
            IMG_HEIGHT = 2.5 / 2.54
            
            # ========== CẤU HÌNH TRANG ==========
            for section in doc.sections:
                section.page_width = Inches(8.27)
                section.page_height = Inches(11.69)
                section.top_margin = Inches(0.3)
                section.bottom_margin = Inches(0.3)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            # ========== CÁC HÀM HỖ TRỢ ==========
            
            def set_cell_width(cell, width_inches):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(width_inches * 1440)))
                tcW.set(qn('w:type'), 'dxa')
                existing = tcPr.find(qn('w:tcW'))
                if existing is not None:
                    tcPr.remove(existing)
                tcPr.append(tcW)

            def set_row_height(row, height_inches):
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), str(int(height_inches * 1440)))
                trHeight.set(qn('w:hRule'), 'exact')
                existing = trPr.find(qn('w:trHeight'))
                if existing is not None:
                    trPr.remove(existing)
                trPr.append(trHeight)

            def set_table_width(table, width_inches):
                tbl = table._element
                tblPr = tbl.tblPr
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)
                tblW = OxmlElement('w:tblW')
                tblW.set(qn('w:w'), str(int(width_inches * 1440)))
                tblW.set(qn('w:type'), 'dxa')
                existing = tblPr.find(qn('w:tblW'))
                if existing is not None:
                    tblPr.remove(existing)
                tblPr.append(tblW)

            def set_table_border(table):
                tbl = table._element
                tblPr = tbl.tblPr
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)
                tblBorders = OxmlElement('w:tblBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '12')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tblBorders.append(border)
                existing = tblPr.find(qn('w:tblBorders'))
                if existing is not None:
                    tblPr.remove(existing)
                tblPr.append(tblBorders)

            def remove_table_border(table):
                tbl = table._element
                tblPr = tbl.tblPr
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)
                tblBorders = OxmlElement('w:tblBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tblBorders.append(border)
                existing = tblPr.find(qn('w:tblBorders'))
                if existing is not None:
                    tblPr.remove(existing)
                tblPr.append(tblBorders)

            def set_cell_vertical_align(cell, align="center"):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), align)
                existing = tcPr.find(qn('w:vAlign'))
                if existing is not None:
                    tcPr.remove(existing)
                tcPr.append(vAlign)

            def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                for side, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
                    node = OxmlElement(f'w:{side}')
                    node.set(qn('w:w'), str(value))
                    node.set(qn('w:type'), 'dxa')
                    tcMar.append(node)
                existing = tcPr.find(qn('w:tcMar'))
                if existing is not None:
                    tcPr.remove(existing)
                tcPr.append(tcMar)

            def set_row_page_break_before(row):
                """Set page break before cho row bằng cách set cho cell đầu tiên"""
                first_cell = row.cells[0]
                tc = first_cell._tc
                # Lấy paragraph đầu tiên trong cell
                p = tc.find(qn('w:p'))
                if p is None:
                    p = OxmlElement('w:p')
                    tc.append(p)
                pPr = p.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    p.insert(0, pPr)
                # Thêm pageBreakBefore
                pageBreakBefore = OxmlElement('w:pageBreakBefore')
                existing = pPr.find(qn('w:pageBreakBefore'))
                if existing is not None:
                    pPr.remove(existing)
                pPr.append(pageBreakBefore)

            def clear_paragraph_formatting(p):
                """Xóa formatting thừa của paragraph"""
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0

            # ========== KIỂM TRA PIL ==========
            try:
                from PIL import Image
                import base64
                from io import BytesIO
                can_use_image = True
            except ImportError:
                can_use_image = False

            # ========== TÍNH TOÁN ==========
            total_cards = len(hoc_vien_list)
            num_pages = (total_cards + CARDS_PER_PAGE - 1) // CARDS_PER_PAGE
            total_rows = num_pages * NUM_ROWS

            # ========== XÓA PARAGRAPH MẶC ĐỊNH ==========
            # Xóa paragraph mặc định của document
            for p in doc.paragraphs:
                p._element.getparent().remove(p._element)

            # ========== TẠO 1 TABLE LỚN DUY NHẤT ==========
            main_table = doc.add_table(rows=total_rows, cols=NUM_COLS)
            main_table.autofit = False
            set_table_border(main_table)
            set_table_width(main_table, CARD_WIDTH * NUM_COLS)

            # ========== DUYỆT TỪNG HÀNG ==========
            for row_idx in range(total_rows):
                page_idx = row_idx // NUM_ROWS
                row_in_page = row_idx % NUM_ROWS
                
                row = main_table.rows[row_idx]
                set_row_height(row, CARD_HEIGHT)
                
                # Set page break before cho hàng đầu tiên của mỗi trang (trừ trang đầu)
                if row_in_page == 0 and page_idx > 0:
                    set_row_page_break_before(row)
                
                # Duyệt từng cột trong hàng
                for col_idx in range(NUM_COLS):
                    card_idx = page_idx * CARDS_PER_PAGE + row_in_page * NUM_COLS + col_idx
                    
                    main_cell = row.cells[col_idx]
                    set_cell_width(main_cell, CARD_WIDTH)
                    set_cell_margins(main_cell, 0, 0, 0, 0)
                    
                    # Nếu hết dữ liệu, để ô trống
                    if card_idx >= total_cards:
                        continue
                    
                    sbd, ho_ten, anh_base64, hang_gplx = hoc_vien_list[card_idx]
                    sbd = str(sbd) if sbd else ""
                    ho_ten = str(ho_ten) if ho_ten else ""
                    hang_gplx = str(hang_gplx) if hang_gplx else ""
                    
                    # Xóa nội dung cũ trong cell
                    main_cell._element.clear_content()
                    
                    # ===== INNER TABLE: 2 rows x 2 cols =====
                    inner = main_cell.add_table(rows=2, cols=2)
                    inner.autofit = False
                    remove_table_border(inner)
                    set_table_width(inner, CARD_WIDTH)
                    
                    set_row_height(inner.rows[0], ROW_HEADER)
                    set_row_height(inner.rows[1], ROW_BODY)
                    
                    for r in range(2):
                        set_cell_width(inner.rows[r].cells[0], COL_IMG)
                        set_cell_width(inner.rows[r].cells[1], COL_INFO)
                        set_cell_margins(inner.rows[r].cells[0], 0, 0, 0, 0)
                        set_cell_margins(inner.rows[r].cells[1], 0, 0, 0, 0)

                    # ===== ROW 0: HEADER =====
                    header_cell = inner.rows[0].cells[0].merge(inner.rows[0].cells[1])
                    set_cell_vertical_align(header_cell, "center")
                    
                    p1 = header_cell.paragraphs[0]
                    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p1.paragraph_format.space_before = Pt(10)
                    p1.paragraph_format.space_after = Pt(0)
                    p1.paragraph_format.line_spacing = 1.0
                    run1 = p1.add_run("TRƯỜNG TCN GT & XD VIỆT ÚC")
                    run1.font.name = "Times New Roman"
                    run1.font.size = Pt(9)
                    run1.font.bold = True

                    p2 = header_cell.add_paragraph()
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p1.paragraph_format.space_before = Pt(10)
                    p1.paragraph_format.space_after = Pt(10)
                    p1.paragraph_format.line_spacing = 1.0
                    run2 = p2.add_run("THẺ DỰ THI")
                    run2.font.name = "Times New Roman"
                    run2.font.size = Pt(14)
                    run2.font.bold = True

                    # ===== ROW 1 - CỘT 0: ẢNH + NGÀY THI =====
                    left_cell = inner.rows[1].cells[0]
                    set_cell_vertical_align(left_cell, "center")

                    p_img = left_cell.paragraphs[0]
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    clear_paragraph_formatting(p_img)
                    p_img.paragraph_format.space_after = Pt(2)

                    if can_use_image and anh_base64:
                        try:
                            img_data = base64.b64decode(anh_base64)
                            img = Image.open(BytesIO(img_data))
                            tmp_path = f"temp_{card_idx}.png"
                            img.save(tmp_path)
                            run_img = p_img.add_run()
                            run_img.add_picture(tmp_path, width=Inches(IMG_WIDTH), height=Inches(IMG_HEIGHT))
                            os.remove(tmp_path)
                        except:
                            p_img.add_run("[Ảnh]").font.size = Pt(8)
                    else:
                        p_img.add_run("[Ảnh]").font.size = Pt(8)

                    p_date = left_cell.add_paragraph()
                    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    clear_paragraph_formatting(p_date)
                    run_date = p_date.add_run(f"Ngày thi: {ngay_thi}")
                    run_date.font.name = "Segoe UI"
                    run_date.font.size = Pt(8)
                    run_date.font.italic = True

                    # ===== ROW 1 - CỘT 1: THÔNG TIN =====
                    right_cell = inner.rows[1].cells[1]
                    set_cell_vertical_align(right_cell, "center")

                    p_lbl = right_cell.paragraphs[0]
                    clear_paragraph_formatting(p_lbl)
                    run_lbl = p_lbl.add_run("Họ và Tên:")
                    run_lbl.font.name = "Segoe UI"
                    run_lbl.font.size = Pt(9)

                    p_name = right_cell.add_paragraph()
                    clear_paragraph_formatting(p_name)
                    p_name.paragraph_format.space_before = Pt(7)
                    p_name.paragraph_format.space_after = Pt(10)
                    run_name = p_name.add_run(ho_ten)
                    run_name.font.name = "Segoe UI"
                    run_name.font.size = Pt(11)
                    run_name.font.bold = True

                    p_hs = right_cell.add_paragraph()
                    clear_paragraph_formatting(p_hs)
                    run_h1 = p_hs.add_run("Hạng: ")
                    run_h1.font.name = "Segoe UI"
                    run_h1.font.size = Pt(9)
                    run_h2 = p_hs.add_run(hang_gplx)
                    run_h2.font.name = "Segoe UI"
                    run_h2.font.size = Pt(12)
                    run_h2.font.bold = True
                    
                    p_hs.add_run("    ")
                    
                    run_s1 = p_hs.add_run("SBD: ")
                    run_s1.font.name = "Segoe UI"
                    run_s1.font.size = Pt(9)
                    run_s2 = p_hs.add_run(sbd)
                    run_s2.font.name = "Segoe UI"
                    run_s2.font.size = Pt(15)
                    run_s2.font.bold = True

            doc.save(save_path)
            
            messagebox.showinfo("Thành công", 
                            f"Đã xuất {total_cards} thẻ dự thi\n"
                            f"Kích thước: {CARD_WIDTH}\" x {CARD_HEIGHT}\"\n"
                            f"Bố cục: {NUM_ROWS} hàng x {NUM_COLS} cột ({CARDS_PER_PAGE} thẻ/trang)\n"
                            f"Số trang: {num_pages}\n"
                            f"File: {save_path}")
            
            write_log("General", f"Xuất Word: {total_cards} thẻ -> {save_path}")
            
        except Exception as e:
            error_msg = f"Lỗi xuất Word:\n{str(e)}\n\n{traceback.format_exc()}"
            messagebox.showerror("Lỗi", error_msg)
            write_log("General", error_msg)



    # Thêm hàm view_import_results
    def view_import_results(self):
        """Hiển thị thống kê kết quả nhập dữ liệu thi"""
        conn = self.get_connection()
        if not conn:
            return
        
        cur = conn.cursor()
        
        # Lấy tổng số học viên
        cur.execute("SELECT COUNT(*) FROM NguoiLX_HoSo")
        total = cur.fetchone()[0]
        
        if total == 0:
            messagebox.showinfo("Thông báo", "Chưa có dữ liệu học viên!")
            return
        
        # Định nghĩa các cột cần thống kê
        columns_to_check = [
            ("LYTHUYETKT", "Lý thuyết"),
            ("MOPHONGKT", "Mô phỏng"),
            ("HINHKT", "Hình"),
            ("DUONGKT", "Đường")
        ]
        
        # Thu thập thống kê
        stats = {}
        for db_col, display_name in columns_to_check:
            # Đếm Đạt (DA)
            cur.execute(f"SELECT COUNT(*) FROM NguoiLX_HoSo WHERE {db_col} = 'DA'")
            dat = cur.fetchone()[0]
            
            # Đếm Trượt (RO)
            cur.execute(f"SELECT COUNT(*) FROM NguoiLX_HoSo WHERE {db_col} = 'RO'")
            truot = cur.fetchone()[0]
            
            # Đếm Vắng (NULL, rỗng, hoặc 'VA')
            cur.execute(f"""
                SELECT COUNT(*) FROM NguoiLX_HoSo 
                WHERE {db_col} IS NULL OR {db_col} = '' OR {db_col} = 'VA'
            """)
            vang = cur.fetchone()[0]
            
            # Đếm khác (nếu có giá trị không phải DA, RO, VA, NULL, '')
            cur.execute(f"""
                SELECT COUNT(*) FROM NguoiLX_HoSo 
                WHERE {db_col} IS NOT NULL AND {db_col} != '' 
                AND {db_col} != 'DA' AND {db_col} != 'RO' AND {db_col} != 'VA'
            """)
            khac = cur.fetchone()[0]
            
            stats[display_name] = {
                "db_col": db_col,
                "dat": dat,
                "truot": truot,
                "vang": vang,
                "khac": khac,
                "da_nhap": dat + truot,
                "chua_nhap": vang
            }
        
        # Tạo cửa sổ popup
        result_win = Toplevel(self.root)
        result_win.title("Thống kê Kết quả Nhập Dữ liệu")
        result_win.geometry("750x550")
        result_win.resizable(True, True)
        
        # Tiêu đề
        title_frame = Frame(result_win)
        title_frame.pack(fill=X, padx=10, pady=10)
        
        Label(title_frame, text="📊 THỐNG KÊ KẾT QUẢ NHẬP DỮ LIỆU THI",
            font=("Arial", 16, "bold"), fg="blue").pack()
        Label(title_frame, text=f"Tổng số học viên: {total}",
            font=("Arial", 12)).pack(pady=5)
        
        # Frame chính
        main_frame = Frame(result_win)
        main_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)
        
        # Tạo Treeview hiển thị thống kê
        columns = ("mon_thi", "da_nhap", "dat", "truot", "vang", "ti_le_dat")
        tree = ttk.Treeview(main_frame, columns=columns, show="headings", height=6)
        
        # Cấu hình cột
        tree.heading("mon_thi", text="Môn thi")
        tree.heading("da_nhap", text="Đã nhập")
        tree.heading("dat", text="✅ Đạt")
        tree.heading("truot", text="❌ Trượt")
        tree.heading("vang", text="⬜ Chưa nhập")
        tree.heading("ti_le_dat", text="Tỷ lệ đạt")
        
        tree.column("mon_thi", width=120, anchor="w")
        tree.column("da_nhap", width=100, anchor="center")
        tree.column("dat", width=100, anchor="center")
        tree.column("truot", width=100, anchor="center")
        tree.column("vang", width=120, anchor="center")
        tree.column("ti_le_dat", width=100, anchor="center")
        
        # Thêm dữ liệu vào bảng
        for mon, data in stats.items():
            da_nhap = data["da_nhap"]
            ti_le = f"{(data['dat']/da_nhap*100):.1f}%" if da_nhap > 0 else "N/A"
            
            tree.insert("", "end", values=(
                mon,
                da_nhap,
                data["dat"],
                data["truot"],
                data["vang"],
                ti_le
            ))
        
        tree.pack(fill=X, pady=10)
        
        # ========== BIỂU ĐỒ THANH NGANG (TEXT-BASED) ==========
        chart_frame = ttk.LabelFrame(main_frame, text="📈 Biểu đồ tiến độ nhập liệu", padding="10")
        chart_frame.pack(fill=X, pady=10)
        
        for mon, data in stats.items():
            row_frame = Frame(chart_frame)
            row_frame.pack(fill=X, pady=3)
            
            # Tên môn
            Label(row_frame, text=f"{mon}:", width=12, anchor="w", 
                font=("Arial", 10, "bold")).pack(side=LEFT)
            
            # Thanh progress
            progress_frame = Frame(row_frame, bg="white", relief="sunken", bd=1)
            progress_frame.pack(side=LEFT, fill=X, expand=True, padx=5)
            
            # Tính phần trăm
            pct_dat = (data["dat"] / total * 100) if total > 0 else 0
            pct_truot = (data["truot"] / total * 100) if total > 0 else 0
            pct_vang = (data["vang"] / total * 100) if total > 0 else 0
            
            # Canvas cho thanh progress
            canvas_width = 400
            canvas_height = 22
            canvas = Canvas(progress_frame, width=canvas_width, height=canvas_height, 
                        bg="white", highlightthickness=0)
            canvas.pack(fill=X)
            
            # Vẽ các phần
            x_start = 0
            
            # Phần Đạt (xanh lá)
            if pct_dat > 0:
                w_dat = int(canvas_width * pct_dat / 100)
                canvas.create_rectangle(x_start, 0, x_start + w_dat, canvas_height, 
                                    fill="#4CAF50", outline="")
                if pct_dat > 8:
                    canvas.create_text(x_start + w_dat/2, canvas_height/2, 
                                    text=f"{data['dat']}", fill="white", font=("Arial", 9, "bold"))
                x_start += w_dat
            
            # Phần Trượt (đỏ)
            if pct_truot > 0:
                w_truot = int(canvas_width * pct_truot / 100)
                canvas.create_rectangle(x_start, 0, x_start + w_truot, canvas_height,
                                    fill="#f44336", outline="")
                if pct_truot > 8:
                    canvas.create_text(x_start + w_truot/2, canvas_height/2,
                                    text=f"{data['truot']}", fill="white", font=("Arial", 9, "bold"))
                x_start += w_truot
            
            # Phần Vắng/Chưa nhập (xám)
            if pct_vang > 0:
                w_vang = int(canvas_width * pct_vang / 100)
                canvas.create_rectangle(x_start, 0, x_start + w_vang, canvas_height,
                                    fill="#9E9E9E", outline="")
                if pct_vang > 8:
                    canvas.create_text(x_start + w_vang/2, canvas_height/2,
                                    text=f"{data['vang']}", fill="white", font=("Arial", 9, "bold"))
            
            # Hiển thị phần trăm đã nhập
            pct_da_nhap = pct_dat + pct_truot
            Label(row_frame, text=f"{pct_da_nhap:.0f}%", width=6, anchor="e",
                font=("Arial", 10)).pack(side=LEFT)
        
        # Chú thích
        legend_frame = Frame(chart_frame)
        legend_frame.pack(fill=X, pady=10)
        
        Label(legend_frame, text="Chú thích:", font=("Arial", 9, "bold")).pack(side=LEFT, padx=5)
        
        # Đạt
        Frame(legend_frame, bg="#4CAF50", width=20, height=15).pack(side=LEFT, padx=2)
        Label(legend_frame, text="Đạt", font=("Arial", 9)).pack(side=LEFT, padx=(0, 15))
        
        # Trượt
        Frame(legend_frame, bg="#f44336", width=20, height=15).pack(side=LEFT, padx=2)
        Label(legend_frame, text="Trượt", font=("Arial", 9)).pack(side=LEFT, padx=(0, 15))
        
        # Chưa nhập
        Frame(legend_frame, bg="#9E9E9E", width=20, height=15).pack(side=LEFT, padx=2)
        Label(legend_frame, text="Chưa nhập", font=("Arial", 9)).pack(side=LEFT)
        
        # ========== THỐNG KÊ TỔNG HỢP ==========
        summary_frame = ttk.LabelFrame(main_frame, text="📋 Tổng hợp", padding="10")
        summary_frame.pack(fill=X, pady=10)
        
        # Tính toán tổng
        total_dat = sum(s["dat"] for s in stats.values())
        total_truot = sum(s["truot"] for s in stats.values())
        total_vang = sum(s["vang"] for s in stats.values())
        total_da_nhap = total_dat + total_truot
        
        # Số học viên đạt tất cả các môn
        cur.execute("""
            SELECT COUNT(*) FROM NguoiLX_HoSo 
            WHERE LYTHUYETKT = 'DA' AND MOPHONGKT = 'DA' 
            AND HINHKT = 'DA' AND DUONGKT = 'DA'
        """)
        all_pass = cur.fetchone()[0]
        
        # Số học viên có ít nhất 1 môn trượt
        cur.execute("""
            SELECT COUNT(*) FROM NguoiLX_HoSo 
            WHERE LYTHUYETKT = 'RO' OR MOPHONGKT = 'RO' 
            OR HINHKT = 'RO' OR DUONGKT = 'RO'
        """)
        has_fail = cur.fetchone()[0]
        
        # Số học viên chưa thi môn nào
        cur.execute("""
            SELECT COUNT(*) FROM NguoiLX_HoSo 
            WHERE (LYTHUYETKT IS NULL OR LYTHUYETKT = '' OR LYTHUYETKT = 'VA')
            AND (MOPHONGKT IS NULL OR MOPHONGKT = '' OR MOPHONGKT = 'VA')
            AND (HINHKT IS NULL OR HINHKT = '' OR HINHKT = 'VA')
            AND (DUONGKT IS NULL OR DUONGKT = '' OR DUONGKT = 'VA')
        """)
        no_result = cur.fetchone()[0]
        
        # Hiển thị
        summary_text = f"""
        🎯 Đạt tất cả 4 môn: {all_pass}/{total} học viên ({all_pass/total*100:.1f}%)
        ⚠️ Có ít nhất 1 môn trượt: {has_fail}/{total} học viên ({has_fail/total*100:.1f}%)
        ⬜ Chưa có kết quả môn nào: {no_result}/{total} học viên ({no_result/total*100:.1f}%)
        """
        
        Label(summary_frame, text=summary_text, font=("Consolas", 11), 
            justify="left", anchor="w").pack(fill=X)
        
        # ========== NÚT BẤM ==========
        btn_frame = Frame(result_win)
        btn_frame.pack(fill=X, pady=10)
        
        Button(btn_frame, text="📋 Xem DS đạt tất cả", width=18,
            command=lambda: self.view_filtered_list("all_pass")).pack(side=LEFT, padx=10)
        Button(btn_frame, text="📋 Xem DS có trượt", width=18,
            command=lambda: self.view_filtered_list("has_fail")).pack(side=LEFT, padx=10)
        Button(btn_frame, text="📋 Xem DS chưa thi", width=18,
            command=lambda: self.view_filtered_list("no_result")).pack(side=LEFT, padx=10)
        Button(btn_frame, text="❌ Đóng", width=12,
            command=result_win.destroy, bg="#FFB6C1").pack(side=RIGHT, padx=10)
        
        write_log("General", f"Xem thống kê kết quả: {total} học viên, Đạt tất cả: {all_pass}")


    #Thêm hàm view_filtered_list
    def view_filtered_list(self, filter_type):
        """Hiển thị danh sách học viên theo bộ lọc"""
        conn = self.get_connection()
        if not conn:
            return
        
        cur = conn.cursor()
        
        # Xác định query và tiêu đề
        if filter_type == "all_pass":
            query = """
                SELECT SO_BAO_DANH, HO_VA_TEN, HANG_GPLX, NOI_DUNG_SH, LYTHUYETKT, MOPHONGKT, HINHKT, DUONGKT, KETQUAKT
                FROM NguoiLX_HoSo 
                WHERE LYTHUYETKT = 'DA' AND MOPHONGKT = 'DA' 
                AND HINHKT = 'DA' AND DUONGKT = 'DA'
                ORDER BY SO_BAO_DANH
            """
            title = "📋 DANH SÁCH ĐẠT TẤT CẢ 4 MÔN"
            color = "#4CAF50"
        elif filter_type == "has_fail":
            query = """
                SELECT SO_BAO_DANH, HO_VA_TEN, HANG_GPLX, NOI_DUNG_SH, LYTHUYETKT, MOPHONGKT, HINHKT, DUONGKT, KETQUAKT
                FROM NguoiLX_HoSo 
                WHERE LYTHUYETKT = 'RO' OR MOPHONGKT = 'RO' 
                OR HINHKT = 'RO' OR DUONGKT = 'RO'
                ORDER BY SO_BAO_DANH
            """
            title = "📋 DANH SÁCH CÓ MÔN TRƯỢT"
            color = "#f44336"
        elif filter_type == "no_result":
            query = """
                SELECT SO_BAO_DANH, HO_VA_TEN, HANG_GPLX, NOI_DUNG_SH, LYTHUYETKT, MOPHONGKT, HINHKT, DUONGKT, KETQUAKT 
                FROM NguoiLX_HoSo 
                WHERE (LYTHUYETKT IS NULL OR LYTHUYETKT = '' OR LYTHUYETKT = 'VA')
                AND (MOPHONGKT IS NULL OR MOPHONGKT = '' OR MOPHONGKT = 'VA')
                AND (HINHKT IS NULL OR HINHKT = '' OR HINHKT = 'VA')
                AND (DUONGKT IS NULL OR DUONGKT = '' OR DUONGKT = 'VA')
                ORDER BY SO_BAO_DANH
            """
            title = "📋 DANH SÁCH CHƯA CÓ KẾT QUẢ"
            color = "#9E9E9E"
        else:
            return
        
        cur.execute(query)
        rows = cur.fetchall()
        
        if not rows:
            messagebox.showinfo("Thông báo", "Không có học viên nào thỏa điều kiện!")
            return
        
        # Tạo cửa sổ
        list_win = Toplevel(self.root)
        list_win.title(title)
        list_win.geometry("900x500")
        list_win.resizable(True, True)
        
        # Tiêu đề
        Label(list_win, text=title, font=("Arial", 14, "bold"), fg=color).pack(pady=10)
        Label(list_win, text=f"Tổng số: {len(rows)} học viên", font=("Arial", 11)).pack()
        
        # Frame chứa bảng
        tree_frame = Frame(list_win)
        tree_frame.pack(fill=BOTH, expand=True, padx=10, pady=10)
        
        # Scrollbar
        scroll_y = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
        scroll_x = ttk.Scrollbar(tree_frame, orient=tk.HORIZONTAL)
        
        # Treeview
        columns = ("sbd", "ho_ten", "hang", "noidung", "lt", "mp", "hinh", "duong", "ketqua")
        tree = ttk.Treeview(tree_frame, columns=columns, show="headings",
                        yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)
        
        scroll_y.config(command=tree.yview)
        scroll_x.config(command=tree.xview)
        
        # Cấu hình cột
        tree.heading("sbd", text="SBD")
        tree.heading("ho_ten", text="Họ và Tên")
        tree.heading("hang", text="Hạng")
        tree.heading("noidung", text="Nội Dung")
        tree.heading("lt", text="Lý thuyết")
        tree.heading("mp", text="Mô phỏng")
        tree.heading("hinh", text="Hình")
        tree.heading("duong", text="Đường")
        tree.heading("ketqua", text="Kết Quả")
        
        tree.column("sbd", width=60, anchor="center")
        tree.column("ho_ten", width=200, anchor="w")
        tree.column("hang", width=60, anchor="center")
        tree.column("noidung", width=60, anchor="center")
        tree.column("lt", width=100, anchor="center")
        tree.column("mp", width=100, anchor="center")
        tree.column("hinh", width=100, anchor="center")
        tree.column("duong", width=100, anchor="center")
        tree.column("ketqua", width=60, anchor="center")
        
        # Thêm dữ liệu
        for i, row in enumerate(rows):
            # Chuyển đổi giá trị hiển thị
            display_row = list(row)
            for j in range(3, 7):  # Các cột kết quả
                val = display_row[j]
                if val == "DA":
                    display_row[j] = "✅ Đạt"
                elif val == "RO":
                    display_row[j] = "❌ Trượt"
                elif val == "VA" or val is None or val == "":
                    display_row[j] = "⬜ Chưa thi"
                else:
                    display_row[j] = val
            
            tag = "odd" if i % 2 == 0 else "even"
            tree.insert("", "end", values=display_row, tags=(tag,))
        
        tree.tag_configure("odd", background="#f8f8f8")
        tree.tag_configure("even", background="#ffffff")
        
        tree.pack(side=LEFT, fill=BOTH, expand=True)
        scroll_y.pack(side=RIGHT, fill=Y)
        scroll_x.pack(side=BOTTOM, fill=X)
        
        # Nút bấm
        btn_frame = Frame(list_win)
        btn_frame.pack(fill=X, pady=10)
        
        def export_list():
            """Xuất danh sách ra Excel"""
            save_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx")],
                initialfile=f"DS_{filter_type}.xlsx"
            )
            if save_path:
                df = pd.DataFrame(rows, columns=["SBD", "Họ và Tên", "Hạng", "Nội Dung", "LT", "MP", "Hình", "Đường", "Kết Quả"])
                df.to_excel(save_path, index=False)
                messagebox.showinfo("Thành công", f"Đã xuất {len(rows)} học viên ra:\n{save_path}")
        
        Button(btn_frame, text="📤 Xuất Excel", command=export_list, width=15).pack(side=LEFT, padx=10)
        Button(btn_frame, text="❌ Đóng", command=list_win.destroy, width=15, bg="#FFB6C1").pack(side=RIGHT, padx=10)


    # Thêm hàm export_result_report
    def export_result_report(self):
        """Xuất báo cáo kết quả thi ra Excel"""
        conn = self.get_connection()
        if not conn:
            return
        
        cur = conn.cursor()
        
        # Lấy dữ liệu
        cur.execute("""
            SELECT SO_BAO_DANH, MA_DK, HO_VA_TEN, NGAY_SINH, SO_CMT, HANG_GPLX, NOI_DUNG_SH,
                LYTHUYETKT, MOPHONGKT, HINHKT, DUONGKT, KETQUAKT
            FROM NguoiLX_HoSo 
            ORDER BY SO_BAO_DANH
        """)
        rows = cur.fetchall()
        
        if not rows:
            messagebox.showwarning("Cảnh báo", "Không có dữ liệu để xuất!")
            return
        
        # Chọn nơi lưu
        save_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx")],
            initialfile="BaoCao_KetQua_Thi.xlsx"
        )
        
        if not save_path:
            return
        
        try:
            # Tạo DataFrame
            df = pd.DataFrame(rows, columns=[
                "SBD", "Mã ĐK", "Họ và Tên", "Ngày sinh", "Số CMT", "Hạng", "Nội Dung",
                "Lý thuyết", "Mô phỏng", "Hình", "Đường", "Kết Quả"
            ])
            
            # Thêm cột kết quả tổng hợp
            def get_overall_result(row):
                results = [row["Lý thuyết"], row["Mô phỏng"], row["Hình"], row["Đường"]]
                if all(r == "DA" for r in results):
                    return "ĐẠT"
                elif any(r == "RO" for r in results):
                    return "TRƯỢT"
                else:
                    return "CHƯA ĐỦ"
            
            df["Kết quả"] = df.apply(get_overall_result, axis=1)
            
            # Chuyển đổi mã sang text
            def convert_code(val):
                if val == "DA":
                    return "Đạt"
                elif val == "RO":
                    return "Trượt"
                elif val == "VA" or pd.isna(val) or val == "":
                    return "Chưa thi"
                return val
            
            for col in ["Lý thuyết", "Mô phỏng", "Hình", "Đường"]:
                df[col] = df[col].apply(convert_code)
            
            # Xuất Excel với định dạng
            with pd.ExcelWriter(save_path, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='Kết quả thi', index=False)
                
                # Định dạng
                from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
                from openpyxl.utils import get_column_letter
                
                workbook = writer.book
                worksheet = writer.sheets['Kết quả thi']
                
                # Header
                header_font = Font(bold=True, color="FFFFFF")
                header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                
                for col_num in range(1, len(df.columns) + 1):
                    cell = worksheet.cell(row=1, column=col_num)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                
                # Tô màu kết quả
                for row_num in range(2, len(df) + 2):
                    for col_num in range(7, 11):  # Cột kết quả các môn
                        cell = worksheet.cell(row=row_num, column=col_num)
                        if cell.value == "Đạt":
                            cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
                        elif cell.value == "Trượt":
                            cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
                    
                    # Cột kết quả tổng
                    cell = worksheet.cell(row=row_num, column=11)
                    if cell.value == "ĐẠT":
                        cell.fill = PatternFill(start_color="00B050", end_color="00B050", fill_type="solid")
                        cell.font = Font(bold=True, color="FFFFFF")
                    elif cell.value == "TRƯỢT":
                        cell.fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")
                        cell.font = Font(bold=True, color="FFFFFF")
                
                # Độ rộng cột
                col_widths = [8, 25, 25, 12, 15, 8, 10, 10, 10, 10, 12]
                for i, width in enumerate(col_widths, 1):
                    worksheet.column_dimensions[get_column_letter(i)].width = width
                
                # Đóng băng header
                worksheet.freeze_panes = 'A2'
            
            messagebox.showinfo("Thành công", f"Đã xuất báo cáo {len(rows)} học viên ra:\n{save_path}")
            write_log("General", f"Xuất báo cáo kết quả thi: {len(rows)} học viên -> {save_path}")
            
            if messagebox.askyesno("Mở file", "Bạn có muốn mở file vừa xuất không?"):
                os.startfile(save_path)
                
        except Exception as e:
            error_msg = f"Lỗi xuất báo cáo: {e}\n{traceback.format_exc()}"
            messagebox.showerror("Lỗi", error_msg)
            write_log("General", error_msg)

    # Thêm hàm view_missing_results
    def view_missing_results(self):
        """Hiển thị danh sách học viên chưa có kết quả theo từng môn"""
        conn = self.get_connection()
        if not conn:
            return
        
        cur = conn.cursor()
        
        # Tạo cửa sổ
        missing_win = Toplevel(self.root)
        missing_win.title("Danh sách chưa có kết quả theo môn")
        missing_win.geometry("600x400")
        missing_win.resizable(True, True)
        
        Label(missing_win, text="📋 DANH SÁCH CHƯA CÓ KẾT QUẢ THEO MÔN",
            font=("Arial", 14, "bold"), fg="orange").pack(pady=10)
        
        # Notebook (Tab)
        notebook = ttk.Notebook(missing_win)
        notebook.pack(fill=BOTH, expand=True, padx=10, pady=5)
        
        # Các môn cần kiểm tra
        subjects = [
            ("LYTHUYETKT", "Lý thuyết"),
            ("MOPHONGKT", "Mô phỏng"),
            ("HINHKT", "Hình"),
            ("DUONGKT", "Đường")
        ]
        
        for db_col, display_name in subjects:
            # Lấy danh sách chưa có kết quả
            cur.execute(f"""
                SELECT SO_BAO_DANH, HO_VA_TEN, HANG_GPLX
                FROM NguoiLX_HoSo 
                WHERE {db_col} IS NULL OR {db_col} = '' OR {db_col} = 'VA'
                ORDER BY SO_BAO_DANH
            """)
            rows = cur.fetchall()
            
            # Tạo tab
            tab_frame = Frame(notebook)
            notebook.add(tab_frame, text=f"{display_name} ({len(rows)})")
            
            # Treeview trong tab
            columns = ("sbd", "ho_ten", "hang")
            tree = ttk.Treeview(tab_frame, columns=columns, show="headings", height=12)
            
            tree.heading("sbd", text="SBD")
            tree.heading("ho_ten", text="Họ và Tên")
            tree.heading("hang", text="Hạng")
            
            tree.column("sbd", width=80, anchor="center")
            tree.column("ho_ten", width=300, anchor="w")
            tree.column("hang", width=80, anchor="center")
            
            for i, row in enumerate(rows):
                tag = "odd" if i % 2 == 0 else "even"
                tree.insert("", "end", values=row, tags=(tag,))
            
            tree.tag_configure("odd", background="#f8f8f8")
            tree.tag_configure("even", background="#ffffff")
            
            # Scrollbar
            scroll = ttk.Scrollbar(tab_frame, orient=tk.VERTICAL, command=tree.yview)
            tree.configure(yscrollcommand=scroll.set)
            
            tree.pack(side=LEFT, fill=BOTH, expand=True)
            scroll.pack(side=RIGHT, fill=Y)
        
        # Nút đóng
        Button(missing_win, text="❌ Đóng", command=missing_win.destroy, 
            width=15, bg="#FFB6C1").pack(pady=10)

    #Thêm sau phương thức export_kq_baocao
    def update_exam_result(self, file_type, result_status):
        """Cập nhật kết quả thi cho học viên được chọn"""
        # Lấy item được chọn trong Treeview
        selected = self.data_tree.selection()
        
        if not selected:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn học viên cần cập nhật kết quả!")
            return
        
        # Mapping loại thi sang cột trong database
        column_mapping = {
            "Lý Thuyết": "LYTHUYETKT",
            "Mô Phỏng": "MOPHONGKT", 
            "Hình": "HINHKT",
            "Đường": "DUONGKT"
        }
        
        db_column = column_mapping.get(file_type)
        if not db_column:
            messagebox.showerror("Lỗi", f"Không tìm thấy cấu hình cho loại: {file_type}")
            return
        
        conn = self.get_connection()
        if not conn:
            return
        
        cur = conn.cursor()
        updated_count = 0
        
        try:
            for item in selected:
                values = self.data_tree.item(item, "values")
                if values and len(values) > 0:
                    sbd = values[0]  # Cột SBD
                    
                    # Cập nhật kết quả
                    cur.execute(f"UPDATE NguoiLX_HoSo SET {db_column} = ? WHERE SO_BAO_DANH = ?", 
                            (result_status if result_status else None, sbd))
                    updated_count += 1
            
            conn.commit()
            
            # Hiển thị kết quả
            result_text = "Đạt" if result_status == "DA" else ("Trượt" if result_status == "RO" else "Xóa")
            messagebox.showinfo("Thành công", 
                            f"Đã cập nhật {updated_count} học viên\n"
                            f"Môn: {file_type}\n"
                            f"Kết quả: {result_text}")
            
            write_log("General", f"Cập nhật KQ {file_type}: {updated_count} học viên -> {result_status}")
            self.refresh_data()
            
        except Exception as e:
            error_msg = f"Lỗi cập nhật kết quả: {e}"
            messagebox.showerror("Lỗi", error_msg)
            write_log("General", error_msg)


    # ---------- XUẤT FILE Excel ----------
    def export_to_excel(self):
        """Xuất dữ liệu NguoiLX_HoSo ra file Excel"""
        conn = self.get_connection()
        if not conn:
            return
        
        try:
            cur = conn.cursor()
            
            # Lấy tất cả dữ liệu
            cur.execute("SELECT * FROM NguoiLX_HoSo ORDER BY SO_BAO_DANH")
            rows = cur.fetchall()
            
            if not rows:
                messagebox.showwarning("Cảnh báo", "Không có dữ liệu để xuất!")
                return
            
            # Lấy tên cột
            cur.execute("PRAGMA table_info(NguoiLX_HoSo)")
            columns = [col[1] for col in cur.fetchall()]
            
            # Chọn nơi lưu file
            save_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
                initialfile="DuLieu_HocVien.xlsx"
            )
            
            if not save_path:
                return
            
            # Tạo DataFrame từ dữ liệu
            df = pd.DataFrame(rows, columns=columns)
            
            # Loại bỏ cột ảnh (dữ liệu base64 quá dài)
            if 'ANH_CHAN_DUNG' in df.columns:
                df['ANH_CHAN_DUNG'] = df['ANH_CHAN_DUNG'].apply(
                    lambda x: '[Có ảnh]' if x and len(str(x)) > 100 else x
                )
            
            # Xuất ra Excel với định dạng
            with pd.ExcelWriter(save_path, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='DuLieu_HocVien', index=False)
                
                # Lấy workbook và worksheet để định dạng
                workbook = writer.book
                worksheet = writer.sheets['DuLieu_HocVien']
                
                # Import để định dạng
                from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
                from openpyxl.utils import get_column_letter
                
                # Định dạng header
                header_font = Font(bold=True, color="FFFFFF", size=11)
                header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                
                # Định dạng border
                thin_border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
                
                # Áp dụng định dạng cho header (dòng 1)
                for col_num, column_title in enumerate(columns, 1):
                    cell = worksheet.cell(row=1, column=col_num)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = header_alignment
                    cell.border = thin_border
                
                # Định dạng dữ liệu
                data_alignment = Alignment(vertical="center", wrap_text=False)
                for row_num in range(2, len(rows) + 2):
                    for col_num in range(1, len(columns) + 1):
                        cell = worksheet.cell(row=row_num, column=col_num)
                        cell.border = thin_border
                        cell.alignment = data_alignment
                        
                        # Tô màu xen kẽ
                        if row_num % 2 == 0:
                            cell.fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
                
                # Tự động điều chỉnh độ rộng cột
                column_widths = {
                    'SO_TT': 8,
                    'MA_DK': 28,
                    'HO_TEN_DEM': 15,
                    'TEN': 12,
                    'HO_VA_TEN': 25,
                    'GIOI_TINH': 10,
                    'NGAY_SINH': 12,
                    'MA_QUOC_TICH': 12,
                    'NOI_CT': 30,
                    'NOI_CT_MA_DVHC': 15,
                    'NOI_CT_MA_DVQL': 15,
                    'SO_CMT': 15,
                    'SO_HO_SO': 15,
                    'MA_KY_SH': 15,
                    'SO_BAO_DANH': 12,
                    'MA_CSDT': 12,
                    'MA_TTSH': 12,
                    'MA_SO_GTVT': 12,
                    'GIAY_CNSK': 15,
                    'HANG_GPLX': 10,
                    'SO_GPLX_DA_CO': 15,
                    'HANG_GPLX_DA_CO': 15,
                    'DVQL_GPLX_DACO': 15,
                    'NGAY_HH_GPLX_DACO': 18,
                    'SO_NAM_LAIXE': 12,
                    'SO_KM_ANTOAN': 15,
                    'SO_GIAY_CNTN': 15,
                    'SO_CCN': 12,
                    'NOI_DUNG_SH': 12,
                    'LY_DO_SH': 12,
                    'KET_QUA_SH': 12,
                    'KQ_SH_LYTHUYET': 15,
                    'KQ_SH_MOPHONG': 15,
                    'KQ_SH_HINH': 12,
                    'KQ_SH_DUONG': 12,
                    'GHI_CHU_SH': 20,
                    'ANH_CHAN_DUNG': 12,
                    'NGAY_TT_GPLX_DACO': 18,
                    'MA_KHOA_HOC': 15,
                    'SO_QD_SH': 15,
                    'NGAY_QD_SH': 12,
                    'NGUOI_QD_SH': 15,
                    'CHAT_LUONG_ANH': 15,
                    'LYTHUYETKT': 12,
                    'MOPHONGKT': 12,
                    'HINHKT': 10,
                    'DUONGKT': 10,
                    'KETQUAKT': 10,
                    'TRANGTHAI': 10
                }
                
                for col_num, column_name in enumerate(columns, 1):
                    col_letter = get_column_letter(col_num)
                    width = column_widths.get(column_name, 15)
                    worksheet.column_dimensions[col_letter].width = width
                
                # Đóng băng dòng header
                worksheet.freeze_panes = 'A2'
                
                # Thiết lập chiều cao dòng header
                worksheet.row_dimensions[1].height = 30
            
            # Thông báo thành công
            messagebox.showinfo("Thành công", 
                            f"Đã xuất {len(rows)} hồ sơ ra file:\n{save_path}")
            
            write_log("General", f"Xuất Excel thành công: {len(rows)} hồ sơ -> {save_path}")
            self.status_label.config(text=f"Đã xuất {len(rows)} hồ sơ ra Excel")
            
            # Hỏi có muốn mở file không
            if messagebox.askyesno("Mở file", "Bạn có muốn mở file Excel vừa xuất không?"):
                os.startfile(save_path)
            
        except Exception as e:
            error_msg = f"Lỗi khi xuất Excel:\n{str(e)}\n\n{traceback.format_exc()}"
            messagebox.showerror("Lỗi", error_msg)
            write_log("General", error_msg)

if __name__ == "__main__":
    write_log("General", "--- KHỞI ĐỘNG PHẦN MỀM ---")
    root = tk.Tk()
    app = ExamDataManager(root)
    root.mainloop()
    write_log("General", "--- KẾT THÚC PHẦN MỀM ---")